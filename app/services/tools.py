"""
粮情分析工具集 (T1-T8)

将各个功能封装为独立的工具，供 Agent 调用
"""

from typing import Dict, Any, List, Optional
from datetime import datetime, timedelta
from pathlib import Path
import os
import logging
import warnings
# 忽略 matplotlib 的字体警告，避免干扰控制台输出
warnings.filterwarnings("ignore", message="Glyph.*missing from current font")

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib import dates as mdates
from app.services.wms_client import WMSClient
from app.services.analysis_service import AnalysisService
from app.services.llm_service import LLMService
import numpy as np
import json, time
from app.models.domain import Reading

logger = logging.getLogger(__name__)


class GrainTools:
    """粮情分析工具集"""
    
    def __init__(self):
        self.wms_client = WMSClient()
        self.analysis_service = AnalysisService()
        self.llm_service = LLMService()

    def _parse_dt(self, value: str) -> datetime:
        """解析 interface_schema.md 约定的时间字符串。"""
        s = (value or "").strip()
        # 支持的格式列表（按优先级）
        formats = [
            "%Y-%m-%d %H:%M:%S",  # 标准格式
            "%Y-%m-%dT%H:%M:%S",  # ISO 格式
            "%Y-%m-%d",            # 日期格式
        ]
        for fmt in formats:
            try:
                return datetime.strptime(s, fmt)
            except Exception:
                pass
        # 最后尝试 ISO
        try:
            return datetime.fromisoformat(s)
        except Exception:
            raise ValueError(f"无法解析时间格式: {value}")

    def _grain_temp_values_to_readings(self, house_code: str, check_time: datetime, temp_values: str, indoor_humidity: Optional[float] = None) -> List[Reading]:
        """把 temp_values 字符串解析成 Reading 列表（用于分析服务）。"""
        readings: List[Reading] = []
        parts = [p for p in (temp_values or "").split("|") if p.strip()]
        for p in parts:
            # 格式: "27.7,1,1,1"
            try:
                val_s, layer_s, row_s, col_s = [x.strip() for x in p.split(",")]
                sensor_id = f"{house_code}-T{layer_s}{row_s}{col_s}"
                readings.append(Reading(sensor_id=sensor_id, timestamp=check_time, value=float(val_s), type="temperature"))
            except Exception:
                continue
        if indoor_humidity is not None:
            readings.append(Reading(sensor_id=f"{house_code}-H1", timestamp=check_time, value=float(indoor_humidity), type="humidity"))
        return readings
    
    # T1: 粮库巡检
    def inspection(self, warehouse_ids: Optional[List[str]] = None) -> Dict[str, Any]:
        """
        T1: 粮库巡检
        
        检查指定粮库的所有仓，识别异常点位
        
        Args:
            warehouse_ids: 粮库ID列表，None表示所有粮库
        
        Returns:
            巡检结果，包含异常点位和统计摘要
        """
        # 将空列表视为"未指定"（LLM 有时会传 warehouse_ids: []）
        if not warehouse_ids or "all" in warehouse_ids:
            warehouse_ids = ["1"]  # Mock: 只有1个粮库
        
        all_issues = []
        total_silos = 0
        abnormal_silos = 0
        
        for wh_id in warehouse_ids:
            warehouse = self.wms_client.get_warehouse(wh_id)
            for silo in warehouse.silos:
                total_silos += 1
                
                # C 对齐：优先使用标准 WMS 粮温接口
                end_time = datetime.now()
                start_time = end_time - timedelta(hours=1)
                temps = self.wms_client.get_grain_temperature(silo.id, start_time, end_time)
                if not temps:
                    continue
                max_temp = max(t.max_temp for t in temps)
                if max_temp > 30:
                    abnormal_silos += 1
                    all_issues.append({
                        "warehouse_id": wh_id,
                        "silo_id": silo.id,
                        "issue": f"温度过高: {max_temp:.1f}°C",
                        "severity": "danger"
                    })
                elif max_temp > 28:
                    abnormal_silos += 1
                    all_issues.append({
                        "warehouse_id": wh_id,
                        "silo_id": silo.id,
                        "issue": f"温度偏高: {max_temp:.1f}°C",
                        "severity": "warning"
                    })
        
        result = {
            "total_silos": total_silos,
            "abnormal_silos": abnormal_silos,
            "issues": all_issues,
            "summary": f"共检查{total_silos}个仓，发现{abnormal_silos}个异常"
        }

        return result
    
    # T2: 数据提取
    def extraction(self, silo_id: str, time_range_hours: int = 24, start_time: Optional[datetime] = None, end_time: Optional[datetime] = None) -> Dict[str, Any]:
        """
        T2: 数据提取
        """
        # V008: 简写解析
        house_code = self.wms_client.resolve_house_code(silo_id)
        
        if start_time and end_time:
            query_start = start_time
            query_end = end_time
        else:
            query_end = datetime.now()
            query_start = query_end - timedelta(hours=time_range_hours)
        
        # C 对齐：通过标准接口取数据
        warehouse_info = self.wms_client.get_warehouse_info(house_code)
        grain_temps = self.wms_client.get_grain_temperature(house_code, query_start, query_end)
        gases = self.wms_client.get_gas_concentration(house_code, query_start, query_end)
        
        # 兼容分析服务：把 temp_values 展开成 Reading（但对外返回序列化 dict）
        readings_models: List[Reading] = []
        for gt in grain_temps:
            try:
                ct = self._parse_dt(gt.check_time)
            except Exception:
                ct = query_end
            readings_models.extend(self._grain_temp_values_to_readings(silo_id, ct, gt.temp_values, gt.indoor_humidity))

        temps_flat = [r.value for r in readings_models if r.type == "temperature"]
        hums_flat = [r.value for r in readings_models if r.type == "humidity"]
        
        stats: Dict[str, Any] = {}
        if temps_flat:
            stats["temperature"] = {"avg": sum(temps_flat) / len(temps_flat), "max": max(temps_flat), "min": min(temps_flat), "count": len(temps_flat)}
        if hums_flat:
            stats["humidity"] = {"avg": sum(hums_flat) / len(hums_flat), "max": max(hums_flat), "min": min(hums_flat), "count": len(hums_flat)}
        
        return {
            "silo_id": silo_id,
            "time_range_hours": time_range_hours,
            "time_range": {"start_time": query_start.strftime("%Y-%m-%d %H:%M:%S"), "end_time": query_end.strftime("%Y-%m-%d %H:%M:%S")},
            "warehouse_info": warehouse_info.model_dump(),
            "grain_temperature": [x.model_dump() for x in grain_temps],
            "gas_concentration": [x.model_dump() for x in gases],
            "total_readings": len(readings_models),
            "readings": [r.model_dump() for r in readings_models],
            "stats": stats,
        }
    
    # T3: 智能分析
    def analysis(self, silo_id: str, readings: Optional[List] = None, start_time: Optional[datetime] = None, end_time: Optional[datetime] = None) -> Dict[str, Any]:
        """
        T3: 智能分析
        """
        if readings is None:
            # 自动提取数据
            extraction_result = self.extraction(silo_id, 24, start_time=start_time, end_time=end_time)
            readings = extraction_result.get("readings", [])

        # readings 可能是 dict（来自 tool 输出），这里统一转为 Reading
        normalized: List[Reading] = []
        for r in (readings or []):
            if isinstance(r, Reading):
                normalized.append(r)
            elif isinstance(r, dict):
                try:
                    normalized.append(Reading(**r))
                except Exception:
                    continue
        readings = normalized

        analysis_result = self.analysis_service.analyze_temperature(silo_id, readings)

        result = {
            "silo_id": analysis_result.silo_id,
            "score": analysis_result.score,
            "findings": analysis_result.findings,
            "analysis_type": analysis_result.analysis_type,
            "risk_level": getattr(analysis_result, "risk_level", None),
        }

        return result

    # T4: 同仓不同时间对比
    def comparison_time(self, silo_id: str, time_windows: List[Dict[str, int]]) -> Dict[str, Any]:
        """
        T4: 同仓不同时间对比
        """
        house_code = self.wms_client.resolve_house_code(silo_id)
        results = []

        for window in time_windows:
            hours_ago = window.get("hours_ago", 24)
            end_time = datetime.now() - timedelta(hours=hours_ago)
            start_time = end_time - timedelta(hours=24)  # 取24小时的数据

            temps = self.wms_client.get_grain_temperature(house_code, start_time, end_time)
            if temps:
                results.append({
                    "window": f"{hours_ago}小时前",
                    "avg_temp": sum(t.avg_temp for t in temps) / len(temps),
                    "max_temp": max(t.max_temp for t in temps),
                    "min_temp": min(t.min_temp for t in temps),
                })

        # 计算差异
        if len(results) >= 2:
            diff = results[0]["avg_temp"] - results[1]["avg_temp"]
            trend = "上升" if diff > 0 else "下降"
        else:
            diff = 0
            trend = "无变化"

        return {
            "silo_id": silo_id,
            "windows": results,
            "temperature_diff": diff,
            "trend": trend,
            "summary": f"温度{trend} {abs(diff):.1f}°C"
        }

    # T5: 不同仓同一时间对比
    def comparison_silo(self, silo_ids: List[str], time_range_hours: int = 24) -> Dict[str, Any]:
        """
        T5: 不同仓同一时间对比
        """
        results = []

        for silo_id in silo_ids:
            house_code = self.wms_client.resolve_house_code(silo_id)
            end_time = datetime.now()
            start_time = end_time - timedelta(hours=time_range_hours)
            temps = self.wms_client.get_grain_temperature(house_code, start_time, end_time)
            if temps:
                results.append({
                    "silo_id": silo_id,
                    "avg_temp": sum(t.avg_temp for t in temps) / len(temps),
                    "max_temp": max(t.max_temp for t in temps),
                    "min_temp": min(t.min_temp for t in temps),
                })
            else:
                results.append({"silo_id": silo_id, "avg_temp": 0, "max_temp": 0, "min_temp": 0})

        # 找出温度最高和最低的仓
        if results:
            highest = max(results, key=lambda x: x["avg_temp"])
            lowest = min(results, key=lambda x: x["avg_temp"])
        else:
            highest = lowest = None

        return {
            "silo_comparison": results,
            "highest_temp_silo": highest,
            "lowest_temp_silo": lowest,
            "summary": f"{highest['silo_id']}号仓温度最高({highest['avg_temp']:.1f}°C)，{lowest['silo_id']}号仓最低({lowest['avg_temp']:.1f}°C)" if highest else "无数据"
        }

    # T6: LLM 推理
    def llm_reasoning(self, query: str, context: Dict[str, Any]) -> Dict[str, Any]:
        """
        T6: LLM 推理

        基于上下文生成智能建议

        Args:
            query: 用户查询
            context: 上下文数据（分析结果、对比结果等）

        Returns:
            LLM 生成的建议和推理过程
        """
        llm_result = self.llm_service.analyze(query, context)

        return {
            "response": llm_result.response,
            "reasoning": llm_result.reasoning
        }

    # === C: WMS 标准数据接口（对齐 interface_schema.md 的 function definitions）===
    def get_warehouse_info(self, house_code: str) -> Dict[str, Any]:
        # V008: 增加简写解析逻辑 (Q1 -> 91620...)
        real_code = self.wms_client.resolve_house_code(house_code)
        info = self.wms_client.get_warehouse_info(real_code)
        return info.model_dump()

    def get_connected_silos(self) -> List[Dict[str, str]]:
        """
        获取当前所有接入智能体的仓房列表。
        """
        return self.wms_client.get_connected_silos()

    def get_grain_temperature(self, house_code: str, start_time: str, end_time: str) -> Dict[str, Any]:
        # V008: 增加简写解析逻辑
        real_code = self.wms_client.resolve_house_code(house_code)
        st = self._parse_dt(start_time)
        et = self._parse_dt(end_time)
        data = self.wms_client.get_grain_temperature(real_code, st, et)
        return {"house_code": real_code, "start_time": start_time, "end_time": end_time, "data": [x.model_dump() for x in data]}

    def get_gas_concentration(self, house_code: str, start_time: str, end_time: str) -> Dict[str, Any]:
        # V008: 增加简写解析逻辑
        real_code = self.wms_client.resolve_house_code(house_code)
        st = self._parse_dt(start_time)
        et = self._parse_dt(end_time)
        data = self.wms_client.get_gas_concentration(real_code, st, et)
        return {"house_code": real_code, "start_time": start_time, "end_time": end_time, "data": [x.model_dump() for x in data]}

    # T7: 可视化
    def visualization(self, silo_id: str, chart_type: str = "line", time_range_hours: Optional[int] = None, start_time: Optional[Any] = None, end_time: Optional[Any] = None) -> Dict[str, Any]:
        """
        T7: 可视化
        """
        house_code = self.wms_client.resolve_house_code(silo_id)
        
        # 配置中文字体，确保图表文字显示正常
        plt.rcParams['font.sans-serif'] = ['SimHei', 'Microsoft YaHei', 'SimSun', 'STHeiti', 'Arial Unicode MS', 'DejaVu Sans']
        plt.rcParams['axes.unicode_minus'] = False

        # 确定时间范围 (处理 LLM 传入的字符串)
        query_start_time = None
        query_end_time = None

        if start_time and end_time:
            try:
                query_start_time = self._parse_dt(start_time) if isinstance(start_time, str) else start_time
                query_end_time = self._parse_dt(end_time) if isinstance(end_time, str) else end_time
            except Exception:
                pass

        if not query_start_time or not query_end_time:
            if time_range_hours is not None:
                query_end_time = datetime.now()
                query_start_time = query_end_time - timedelta(hours=time_range_hours)
            else:
                query_end_time = datetime.now()
                query_start_time = query_end_time - timedelta(hours=24)
        
        series = self.wms_client.get_grain_temperature(house_code, query_start_time, query_end_time)
        
        # V008: 增加排序逻辑，确保 7 个点的选取是准确的
        if series:
            series = sorted(series, key=lambda x: x.check_time)
        
        # V008: 落实规范 - 数据不足熔断 (<= 1 无法绘图)
        if not series or len(series) <= 1:
            msg = f"当前该仓房({silo_id})仅有 {len(series) if series else 0} 次检测记录，无法绘制趋势图，请检查数据采集状态。"
            logger.info(msg)
            return {"silo_id": silo_id, "chart_type": chart_type, "status": "no_data", "message": msg}

        # V008: 增加 7 个数据点的限制 (与三温两湿图对齐)
        if len(series) > 7:
            if start_time is not None:
                series = series[:7] # 指定时间段取前7个
            else:
                series = series[-7:] # 否则取最近7个

        out_dir = Path(__file__).resolve().parents[2] / "artifacts" / "charts"
        out_dir.mkdir(parents=True, exist_ok=True)
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        out_path = out_dir / f"{silo_id}_{chart_type}_{ts}.png"

        if not series:
            plt.figure(figsize=(8, 3))
            plt.title(f"{silo_id} (no data)")
            plt.savefig(out_path, dpi=150, bbox_inches="tight")
            plt.close()
            return {"silo_id": silo_id, "chart_type": chart_type, "file_path": str(out_path), "status": "no_data"}

        if chart_type in ("heatmap", "3d"):
            # ... (保持 heatmap 逻辑)
            last = series[-1]
            grid: Dict[tuple[int, int], List[float]] = {}
            for p in (last.temp_values or "").split("|"):
                if not p.strip(): continue
                try:
                    val_s, layer_s, row_s, col_s = [x.strip() for x in p.split(",")]
                    key = (int(row_s), int(col_s))
                    grid.setdefault(key, []).append(float(val_s))
                except Exception: continue
            mat = [[0.0 for _ in range(3)] for _ in range(3)]
            for r in range(1, 4):
                for c in range(1, 4):
                    vals = grid.get((r, c), [])
                    mat[r - 1][c - 1] = sum(vals) / len(vals) if vals else 0.0
            plt.figure(figsize=(4, 3))
            plt.imshow(mat, cmap="hot", aspect="auto")
            plt.colorbar(label="°C")
            plt.title(f"{silo_id} heatmap ({last.check_time})")
            plt.savefig(out_path, dpi=150, bbox_inches="tight")
            plt.close()
        else:
            # line 逻辑
            times = []
            avg_temps = []
            for record in series:
                try:
                    check_time = self._parse_dt(record.check_time)
                    times.append(check_time)
                    avg_temps.append(record.avg_temp)
                except Exception: continue
            
            plt.figure(figsize=(8, 3))
            plt.plot(times, avg_temps, marker="o", linewidth=2.0, markersize=4)
            plt.xlabel("检测日期", fontsize=10)
            plt.ylabel("温度 (℃)", fontsize=10)
            plt.title(f"{silo_id} 平均粮温趋势 (限7点)", fontsize=12)
            plt.gca().xaxis.set_major_formatter(mdates.DateFormatter("%Y-%m-%d"))
            plt.xticks(rotation=30, ha="right")
            plt.tight_layout()
            plt.savefig(out_path, dpi=150, bbox_inches="tight")
            plt.close()

        return {"silo_id": silo_id, "chart_type": chart_type, "file_path": str(out_path), "status": "generated"}

    # V008 新增：独立工具封装
    def three_temp_chart(self, silo_id: str, start_time: Optional[str] = None, end_time: Optional[str] = None) -> Dict[str, Any]:
        """T10: 独立三温图工具"""
        st = self._parse_dt(start_time) if start_time else None
        et = self._parse_dt(end_time) if end_time else None
        return self.generate_three_temp_chart(silo_id, start_time=st, end_time=et)

    def two_humidity_chart(self, silo_id: str, start_time: Optional[str] = None, end_time: Optional[str] = None) -> Dict[str, Any]:
        """T11: 独立两湿图工具"""
        st = self._parse_dt(start_time) if start_time else None
        et = self._parse_dt(end_time) if end_time else None
        return self.generate_two_humidity_chart(silo_id, start_time=st, end_time=et)

    # T9: 粮情短期预测 (Holt 指数平滑 + 加权线性回归 集成)
    def short_term_prediction(
        self,
        silo_id: str,
        prediction_days: int = 3,
        start_time: Optional[str] = None,
        end_time: Optional[str] = None,
        include_spatial: bool = False,  # P1: 是否包含空间热点分析
    ) -> Dict[str, Any]:
        """
        T9: 粮情短期预测

        基于 Holt 双参数指数平滑（阻尼趋势）与加权线性回归的 Ensemble 预测。
        同时预测 avg_temp / max_temp / min_temp 三个指标，附带 95% 置信区间
        和多因素风险评分。

        返回值完全向后兼容旧字段（predicted_avg_temp / trend / risk_assessment / context），
        同时附加 confidence_interval / daily_predictions / multi_metric / data_quality / risk_details。
        """

        # ================================================================
        # 局部嵌套函数 —— 仅此方法内部可见，不污染 GrainTools 类接口
        # ================================================================

        def detect_sensor_faults(vals: np.ndarray) -> np.ndarray:
            """粮储领域感知的传感器故障检测。

            只剔除物理不可能的读数（传感器故障），保留所有合理范围内的值。
            高温读数可能是真实的热点萌发信号，不应被统计方法剔除。

            判定规则:
              - 物理不可能值: < -40°C 或 > 60°C
              - 相对中位数跳变 > 15°C（传感器瞬时故障特征）
            """
            mask = np.ones(len(vals), dtype=bool)
            median = float(np.median(vals))
            for i, v in enumerate(vals):
                if v < -40 or v > 60:
                    mask[i] = False
                elif abs(v - median) > 15:
                    mask[i] = False
            return mask

        def holt_linear(
            vals: np.ndarray, alpha: float, beta: float, horizon: int, phi: float = 0.98
        ):
            """
            Holt 双参数指数平滑（带阻尼因子 φ）。

            Level:    l_t = α·y_t + (1-α)·(l_{t-1} + φ·b_{t-1})
            Trend:    b_t = β·(l_t - l_{t-1}) + (1-β)·φ·b_{t-1}
            Forecast: ŷ_{t+h} = l_t + Σ(φ^j, j=1..h)·b_t

            Returns: (fitted, forecasts)
            """
            n = len(vals)
            level = np.zeros(n)
            trend = np.zeros(n)
            fitted = np.zeros(n)

            level[0] = vals[0]
            trend[0] = (vals[min(n - 1, 1)] - vals[0]) if n > 1 else 0.0
            fitted[0] = vals[0]

            for t in range(1, n):
                level[t] = alpha * vals[t] + (1 - alpha) * (
                    level[t - 1] + phi * trend[t - 1]
                )
                trend[t] = (
                    beta * (level[t] - level[t - 1]) + (1 - beta) * phi * trend[t - 1]
                )
                fitted[t] = level[t - 1] + phi * trend[t - 1]

            # 阻尼累积系数: Σ φ^j (j=1..h)
            forecasts = np.array(
                [
                    level[n - 1] + sum(phi**j for j in range(1, h + 1)) * trend[n - 1]
                    for h in range(1, horizon + 1)
                ]
            )
            return fitted, forecasts

        def phi_for_month(month: int) -> float:
            """
            季节自适应阻尼因子 φ（方案A）。
            冬季趋势外推容易过头，用更强的阻尼防止方向性误判：
              冬季(12,1,2)：φ=0.90  过渡季(10,11,3,4)：φ=0.94  其余：φ=0.98
            """
            if month in (12, 1, 2):
                return 0.90
            if month in (10, 11, 3, 4):
                return 0.94
            return 0.98

        def optimize_holt(vals: np.ndarray, last_month: int = 0):
            """
            网格搜索最优 (α, β)，最小化单步拟合 RMSE。

            方案B — 冬季偏向网格：
              冬季(12,1,2): α ∈ [0.4,0.9]，β ∈ [0.05,0.15]
                高α = 多信任当前水平，低β = 少跟随趋势
              其他月份: α ∈ [0.1,0.9]，β ∈ [0.05,0.50]（原搜索空间）
            """
            if last_month in (12, 1, 2):
                alpha_range = np.arange(0.4, 1.0, 0.1)
                beta_range = np.arange(0.05, 0.20, 0.05)
            else:
                alpha_range = np.arange(0.1, 1.0, 0.1)
                beta_range = np.arange(0.05, 0.55, 0.05)
            best_rmse = float("inf")
            best_params = (0.3, 0.1)
            for a in alpha_range:
                for b in beta_range:
                    fitted, _ = holt_linear(vals, float(a), float(b), 1)
                    err = vals[1:] - fitted[1:]
                    rmse = float(np.sqrt(np.mean(err**2)))
                    if rmse < best_rmse:
                        best_rmse = rmse
                        best_params = (round(float(a), 2), round(float(b), 2))
            return best_params

        def weighted_regression(
            vals: np.ndarray,
            horizon: int,
            decay: float = 0.15,
            t_days: Optional[np.ndarray] = None,
        ):
            """
            指数衰减加权线性回归。

            权重: w_i = exp(-λ·(T_max - t_i))  —— 越近期权重越大
            使用 numpy.polyfit(deg=1) 拟合直线，外推 horizon 天。

            参数:
              t_days: 各数据点的真实时间轴（以天为单位），如未提供则用等间距序号。

            Returns: (forecasts, residuals)
            """
            n = len(vals)
            if t_days is not None and len(t_days) == n:
                t = t_days.copy()
            else:
                t = np.arange(n, dtype=float)
            weights = np.exp(-decay * (t[-1] - t))
            coeffs = np.polyfit(t, vals, deg=1, w=np.sqrt(weights))
            fitted = np.polyval(coeffs, t)
            residuals = vals - fitted
            # 外推: 从最后一个点开始，每天一个预测点
            last_t = t[-1]
            future_t = np.array(
                [last_t + d for d in range(1, horizon + 1)], dtype=float
            )
            forecasts = np.polyval(coeffs, future_t)
            return forecasts, residuals

        def compute_ci_half(
            residuals: np.ndarray,
            horizon: int,
            n: int,
            n_effective: int = 0,
            last_month: int = 0,
            sigma_floor: float = 0.0,
        ) -> float:
            """
            自适应 95% 预测区间半宽。

            三个膨胀因子叠加：
              A. t 分布小样本补偿 — 样本少时用 t 临界值代替固定 z=1.96
              B. 季节性膨胀       — 秋冬（9-2月）波动大，×1.4
              C. 波动性膨胀       — 残差变异系数大时放宽

            sigma_floor: 预测误差标准差下界（由 L3 留出集估算传入）。
              当拟合残差远小于真实预测误差时（典型于 STL 大数据量场景），
              用 max(拟合残差 sigma, sigma_floor) 防止 CI 过度收窄。
            """
            if len(residuals) <= 1:
                return 0.0
            sigma = max(float(np.std(residuals, ddof=1)), sigma_floor)

            # A. t 分布小样本补偿（自由度 = n_eff - 1）
            n_eff = max(n_effective if n_effective > 0 else len(residuals), 2)
            if n_eff < 5:
                t_val = 4.30
            elif n_eff < 10:
                t_val = 2.26
            elif n_eff < 20:
                t_val = 2.09
            else:
                t_val = 1.96

            # B. 季节性膨胀 — 秋冬季预测难度高
            season_factor = 1.4 if last_month in (9, 10, 11, 12, 1, 2) else 1.0

            # C. 波动性膨胀 — 残差离散程度
            mean_abs = float(np.mean(np.abs(residuals))) or 1e-6
            cv_resid = sigma / mean_abs
            if cv_resid > 0.5:
                vol_factor = 1.3
            elif cv_resid > 0.3:
                vol_factor = 1.15
            else:
                vol_factor = 1.0

            z_adaptive = t_val * season_factor * vol_factor
            return z_adaptive * sigma * float(np.sqrt(1 + horizon / max(n, 1)))

        def assess_risk(
            pred_avg: float,
            pred_max: float,
            daily_rate: float,
            cv: float,
            ci_half: float,
            horizon: int,
            n_pts: int,
            env_ctx: Optional[Dict[str, Any]] = None,
        ) -> Dict[str, Any]:
            """
            多因素风险评分（0-100，越高越安全）。

            6 个扣分因子:
              1. 预测均温绝对值 vs 安全阈值 (25/28/30°C)   最多 -35
              2. 日均变化速率                                最多 -25
              3. 预测最高温 vs 危险阈值                      最多 -20
              4. 历史数据变异系数 (CV)                       最多 -10
              5. 置信区间宽度 + 外推比例                     最多 -15
              6. 环境趋势分歧（气温升粮温未升）              最多 -10

            映射: >=80 low / >=60 medium / >=40 high / <40 critical
            """
            score = 100.0
            factors: List[str] = []

            # 因子 1: 预测均温绝对值
            if pred_avg >= 30:
                score -= 35
                factors.append(f"预测均温 {pred_avg:.1f}°C 达到危险阈值 30°C")
            elif pred_avg >= 28:
                score -= 20
                factors.append(f"预测均温 {pred_avg:.1f}°C 超过警戒线 28°C")
            elif pred_avg >= 25:
                score -= 10
                factors.append(f"预测均温 {pred_avg:.1f}°C 超过安全上限 25°C")
            else:
                factors.append(f"预测均温 {pred_avg:.1f}°C 低于安全阈值 25°C")

            # 因子 2: 日均变化速率
            abs_rate = abs(daily_rate)
            if abs_rate > 0.5:
                score -= 25
                factors.append(f"日均变化 {daily_rate:+.2f}°C/天，变化剧烈")
            elif abs_rate > 0.2:
                score -= 15
                factors.append(f"日均变化 {daily_rate:+.2f}°C/天，变化较快")
            elif abs_rate > 0.05:
                score -= 5
                factors.append(f"日均变化 {daily_rate:+.2f}°C/天，温和变化")
            else:
                factors.append(f"日均变化 {daily_rate:+.2f}°C/天，基本稳定")

            # 因子 3: 预测最高温
            if pred_max >= 30:
                score -= 20
                factors.append(f"预测最高温 {pred_max:.1f}°C 达危险阈值")
            elif pred_max >= 28:
                score -= 10
                factors.append(f"预测最高温 {pred_max:.1f}°C 接近警戒线")
            else:
                factors.append(f"预测最高温 {pred_max:.1f}°C 在安全范围内")

            # 因子 4: 变异系数
            if cv > 0.15:
                score -= 10
                factors.append(f"历史数据波动较大 (CV={cv:.2f})，预测不确定性高")
            elif cv > 0.05:
                score -= 5
                factors.append(f"历史数据有一定波动 (CV={cv:.2f})")
            else:
                factors.append(f"历史数据较为稳定 (CV={cv:.2f})")

            # 因子 5: 置信区间宽度 + 外推比例
            if ci_half > 5:
                score -= 10
                factors.append(f"置信区间较宽 (±{ci_half:.1f}°C)，预测可靠性有限")
            elif ci_half > 2:
                score -= 5
                factors.append(f"置信区间适中 (±{ci_half:.1f}°C)")
            else:
                factors.append(f"置信区间较窄 (±{ci_half:.1f}°C)，预测较可靠")

            if horizon > n_pts * 0.5:
                score -= 5
                factors.append(
                    f"外推天数 ({horizon}天) 超过有效数据量的 50%，结论仅供参考"
                )

            # 因子 6: 环境趋势分歧（气温→粮温的滞后关联）
            if env_ctx:
                if env_ctx.get("warning"):
                    score -= 10
                    factors.append(env_ctx["warning"])
                elif env_ctx.get("info"):
                    factors.append(env_ctx["info"])

            score = max(0.0, min(100.0, score))
            if score >= 80:
                level = "low"
            elif score >= 60:
                level = "medium"
            elif score >= 40:
                level = "high"
            else:
                level = "critical"

            return {
                "risk_level": level,
                "risk_score": round(score, 1),
                "factors": factors,
            }

        def predict_one_metric(
            vals: np.ndarray,
            horizon: int,
            t_days: Optional[np.ndarray] = None,
            last_month: int = 0,
        ):
            """Layer 1 (3-14条): Holt + WLR ensemble 预测。"""
            alpha, beta = optimize_holt(vals, last_month)
            _, holt_fc = holt_linear(
                vals, alpha, beta, horizon, phi=phi_for_month(last_month)
            )
            wlr_fc, residuals = weighted_regression(vals, horizon, t_days=t_days)
            ensemble_fc = 0.6 * holt_fc + 0.4 * wlr_fc
            ci_half = compute_ci_half(
                residuals,
                horizon,
                len(vals),
                n_effective=len(vals),
                last_month=last_month,
            )
            rate = float((ensemble_fc[-1] - vals[-1]) / horizon) if horizon > 0 else 0.0
            return ensemble_fc, ci_half, rate

        # ────────── Layer 2: Holt-Winters 季节性平滑 (15-49条) ──────────

        def detect_period(vals: np.ndarray, t_days: Optional[np.ndarray] = None) -> int:
            """自动检测采样周期 m（每日数据点数）。"""
            if t_days is not None and len(t_days) >= 2:
                intervals = np.diff(t_days)
                median_interval = float(np.median(intervals))
                if median_interval > 0:
                    m = max(2, int(round(1.0 / median_interval)))
                    return min(m, len(vals) // 3)  # 周期不超过数据量的 1/3
            return 2  # 默认

        def holt_winters(
            vals: np.ndarray,
            m: int,
            horizon: int,
            alpha: float = 0.3,
            beta: float = 0.1,
            gamma: float = 0.3,
        ):
            """
            Holt-Winters 加法季节性模型。

            Level:    l_t = α·(y_t - s_{t-m}) + (1-α)·(l_{t-1} + b_{t-1})
            Trend:    b_t = β·(l_t - l_{t-1}) + (1-β)·b_{t-1}
            Seasonal: s_t = γ·(y_t - l_t) + (1-γ)·s_{t-m}
            Forecast: ŷ_{t+h} = l_t + h·b_t + s_{t-m+(h mod m)}

            Returns: (fitted, forecasts, residuals)
            """
            n = len(vals)
            level = np.zeros(n)
            trend = np.zeros(n)
            season = np.zeros(n + horizon)
            fitted = np.zeros(n)

            # 初始化: 第一个周期的均值做 level，趋势用前两周期差
            level[0] = float(np.mean(vals[:m]))
            trend[0] = (
                float(np.mean(vals[m : 2 * m]) - np.mean(vals[:m])) / m
                if n >= 2 * m
                else 0.0
            )
            for j in range(m):
                season[j] = vals[j] - level[0] if j < n else 0.0
            fitted[0] = level[0] + season[0]

            for t in range(1, n):
                s_prev = season[t - m] if t >= m else season[t % m]
                level[t] = alpha * (vals[t] - s_prev) + (1 - alpha) * (
                    level[t - 1] + trend[t - 1]
                )
                trend[t] = beta * (level[t] - level[t - 1]) + (1 - beta) * trend[t - 1]
                season[t] = gamma * (vals[t] - level[t]) + (1 - gamma) * s_prev
                fitted[t] = level[t - 1] + trend[t - 1] + s_prev

            residuals = vals - fitted
            forecasts = np.array(
                [
                    level[n - 1]
                    + (h + 1) * trend[n - 1]
                    + season[n - m + ((h + 1) % m)]
                    for h in range(horizon)
                ]
            )
            return fitted, forecasts, residuals

        def optimize_hw(vals: np.ndarray, m: int):
            """网格搜索最优 Holt-Winters 参数 (α, β, γ)。"""
            best_rmse = float("inf")
            best_params = (0.3, 0.1, 0.3)
            for a in [0.1, 0.3, 0.5, 0.7]:
                for b in [0.05, 0.1, 0.2]:
                    for g in [0.1, 0.3, 0.5]:
                        try:
                            fitted, _, _ = holt_winters(vals, m, 1, a, b, g)
                            err = vals[m:] - fitted[m:]  # 跳过初始化阶段
                            if len(err) > 0:
                                rmse = float(np.sqrt(np.mean(err**2)))
                                if rmse < best_rmse:
                                    best_rmse = rmse
                                    best_params = (a, b, g)
                        except Exception:
                            continue
            return best_params

        def predict_one_metric_hw(
            vals: np.ndarray,
            horizon: int,
            t_days: Optional[np.ndarray] = None,
            last_month: int = 0,
        ):
            """Layer 2 (15-49条): Holt-Winters 季节性预测。"""
            m = detect_period(vals, t_days)
            if m < 2 or len(vals) < 2 * m:
                return predict_one_metric(vals, horizon, t_days, last_month)

            # 周期性检验 — 无明显周期时退化到 Layer 1
            win = min(m, len(vals))
            detrended = vals - np.convolve(vals, np.ones(win) / win, mode="same")
            if len(detrended) > m:
                corr = float(np.corrcoef(detrended[:-m], detrended[m:])[0, 1])
            else:
                corr = 0.0
            if abs(corr) < 0.3:
                return predict_one_metric(vals, horizon, t_days, last_month)

            a, b, g = optimize_hw(vals, m)
            _, forecasts, residuals = holt_winters(vals, m, horizon, a, b, g)

            # 自适应 CI（含小样本、季节、波动性膨胀）
            effective_resid = residuals[m:]
            ci_half = compute_ci_half(
                effective_resid,
                horizon,
                len(vals),
                n_effective=len(effective_resid),
                last_month=last_month,
            )
            rate = float((forecasts[-1] - vals[-1]) / horizon) if horizon > 0 else 0.0
            return forecasts, ci_half, rate

        # ────────── Layer 3: Holt-Winters + STL 分解 (50-199条) ──────────

        def stl_decompose(vals: np.ndarray, m: int):
            """
            简化版 STL 分解: y = trend + seasonal + residual。

            Trend:    移动平均 (窗口=m)
            Seasonal: 去趋势后按周期折叠取中位数
            Residual: y - trend - seasonal
            """
            n = len(vals)
            # Trend: 居中移动平均
            trend = np.full(n, np.nan)
            half = m // 2
            for i in range(half, n - half):
                trend[i] = float(np.mean(vals[max(0, i - half) : i + half + 1]))
            # 填补两端 NaN
            first_valid = half
            last_valid = n - half - 1
            trend[:first_valid] = trend[first_valid]
            trend[last_valid + 1 :] = trend[last_valid]

            # Seasonal: 去趋势 → 按周期位置折叠 → 取中位数
            detrended = vals - trend
            seasonal = np.zeros(n)
            for j in range(m):
                indices = list(range(j, n, m))
                seasonal_val = float(np.median(detrended[indices]))
                for idx in indices:
                    seasonal[idx] = seasonal_val

            residual = vals - trend - seasonal
            return trend, seasonal, residual

        def predict_one_metric_stl(
            vals: np.ndarray,
            horizon: int,
            t_days: Optional[np.ndarray] = None,
            last_month: int = 0,
        ):
            """Layer 3 (50+条): STL 分解 + Holt 趋势外推。"""
            m = detect_period(vals, t_days)
            if m < 2:
                m = max(2, len(vals) // 14)  # 回退: 按 ~14天一个周期

            trend, seasonal, residual = stl_decompose(vals, m)

            # 对 trend 分量用 Holt 外推（传入季节参数，应用A+B改进）
            alpha, beta = optimize_holt(trend, last_month)
            _, trend_fc = holt_linear(
                trend, alpha, beta, horizon, phi=phi_for_month(last_month)
            )

            # seasonal 分量: 用最近一个完整周期的模式循环复制
            n = len(vals)
            seasonal_fc = np.array([seasonal[n - m + (h % m)] for h in range(horizon)])

            forecasts = trend_fc + seasonal_fc

            # ── OOS sigma 估算（留出集）────────────────────────────────────
            # STL 拟合残差在数据量大时极小（模型贴合历史好），但预测误差
            # 会因趋势外推和季节模式漂移而远大于拟合残差。
            # 用最后 k 个点做留出验证，得到真实预测误差标准差作为下界。
            k = max(3, min(15, n // 5))  # 留出最后 20%，min=3，max=15
            sigma_floor = 0.0
            if n - k >= 50:  # 保证训练集至少剩 50 个点
                try:
                    vals_tr = vals[:-k]
                    t_tr = t_days[:-k] if t_days is not None else None
                    m_h = max(2, detect_period(vals_tr, t_tr))
                    t_h, s_h, _ = stl_decompose(vals_tr, m_h)
                    a_h, b_h = optimize_holt(t_h, last_month)
                    _, tfc_h = holt_linear(
                        t_h, a_h, b_h, k, phi=phi_for_month(last_month)
                    )
                    sfc_h = np.array(
                        [s_h[len(vals_tr) - m_h + (j % m_h)] for j in range(k)]
                    )
                    oos_err = vals[-k:] - (tfc_h + sfc_h)
                    if len(oos_err) > 1:
                        sigma_floor = float(np.std(oos_err, ddof=1))
                except Exception:
                    sigma_floor = 0.0  # 留出计算失败时静默降级，不影响预测值
            # ────────────────────────────────────────────────────────────────

            # 自适应 CI（传入 sigma_floor 防止 CI 过度收窄）
            ci_half = compute_ci_half(
                residual,
                horizon,
                len(vals),
                n_effective=len(vals),
                last_month=last_month,
                sigma_floor=sigma_floor,
            )
            rate = float((forecasts[-1] - vals[-1]) / horizon) if horizon > 0 else 0.0
            return forecasts, ci_half, rate

        # ================================================================
        # 主流程
        # ================================================================

        def _build_data_quality_note(
            n_pts: int, clamped: bool, experimental: bool
        ) -> str:
            """构建 data_quality.note 字段，合并多种警告。"""
            notes = []

            # P0-2: 试验性预测警告
            if experimental:
                notes.append(
                    f"⚠️ 试验性预测：当前基于 {n_pts} 个数据点（< 14），"
                    f"预测结果不确定性极高，仅供参考，"
                    f"强烈建议结合现场巡检综合判断。"
                )
            elif n_pts < 30:
                notes.append(
                    f"当前基于 {n_pts} 个数据点的统计预测，数据量有限，"
                    f"建议结合现场巡检综合判断。"
                )

            # 物理下界截断警告
            if clamped:
                notes.append(
                    "预测值触达物理下界（-10°C）已自动截断，"
                    "这通常发生在季节转折点的过度外推，请结合实际情况判断。"
                )

            return " ".join(notes)

        # 1. 获取数据（与原逻辑一致）
        house_code = self.wms_client.resolve_house_code(silo_id)
        if start_time and end_time:
            ref_start = self._parse_dt(start_time)
            ref_end = self._parse_dt(end_time)
        else:
            ref_end = datetime.now()
            ref_start = ref_end - timedelta(days=14)

        series = self.wms_client.get_grain_temperature(house_code, ref_start, ref_end)

        # 2. P0-2: 分层硬阈值检查数据充足性
        from app.services.predictor.data_validator import DataValidator

        n_raw = len(series) if series else 0
        is_sufficient, sufficiency_msg = DataValidator.check_data_sufficiency(n_raw)

        if not is_sufficient:
            logger.info(f"数据不足熔断: {sufficiency_msg}")
            return {
                "silo_id": silo_id,
                "status": "insufficient_data",
                "message": sufficiency_msg,
                "min_required": 7,
                "actual_count": n_raw,
            }

        # 3. 预处理：按时间排序 → 提取三指标 → 传感器故障检测
        series_sorted = sorted(series, key=lambda s: s.check_time)
        raw_avg = np.array([s.avg_temp for s in series_sorted], dtype=float)
        raw_max = np.array([s.max_temp for s in series_sorted], dtype=float)
        raw_min = np.array([s.min_temp for s in series_sorted], dtype=float)

        # 构造真实时间轴（天数），解决非等间隔采样问题
        base_time = datetime.strptime(series_sorted[0].check_time, "%Y-%m-%d %H:%M:%S")
        raw_t_days = np.array(
            [
                (
                    datetime.strptime(s.check_time, "%Y-%m-%d %H:%M:%S") - base_time
                ).total_seconds()
                / 86400.0
                for s in series_sorted
            ],
            dtype=float,
        )

        # P0-1: 使用新的 DataValidator 进行异常检测
        mask, anomaly_stats = DataValidator.detect_anomalies(
            raw_avg, raw_t_days, sensitivity="medium"
        )
        avg_clean = raw_avg[mask]
        max_clean = raw_max[mask]
        min_clean = raw_min[mask]
        t_days_clean = raw_t_days[mask]

        total_pts = len(raw_avg)
        outliers_cnt = int(np.sum(~mask))
        effective_pts = len(avg_clean)

        # 去异常值后仍不足（二次检查）
        is_sufficient_after_clean, msg_after_clean = (
            DataValidator.check_data_sufficiency(effective_pts)
        )
        if not is_sufficient_after_clean:
            msg = f"去除 {outliers_cnt} 个异常值后，有效数据仅 {effective_pts} 条。{msg_after_clean}"
            logger.info(msg)
            return {"silo_id": silo_id, "status": "insufficient_data", "message": msg}

        # 4. 根据数据量自动路由到对应算法层级
        #    当某指标全为 0（API 未提供该字段）时标记为不可用
        def is_metric_available(vals: np.ndarray) -> bool:
            """判断指标数据是否有效（非全零且有方差）。"""
            return float(np.max(np.abs(vals))) > 0.01

        # P0-2: 判断是否为试验性预测（7-14 条数据）
        is_experimental = effective_pts < 14

        if effective_pts < 15:
            method_name = "holt_wlr_lite"
            predict_fn = predict_one_metric
        elif effective_pts < 50:
            method_name = "holt_winters_seasonal"
            predict_fn = predict_one_metric_hw
        else:
            method_name = "holt_winters_stl"
            predict_fn = predict_one_metric_stl

        # 提取训练数据末端月份，用于自适应 CI 的季节感知
        last_record_time = datetime.strptime(
            series_sorted[-1].check_time, "%Y-%m-%d %H:%M:%S"
        )
        last_month = last_record_time.month

        logger.info(
            f"预测路由: {effective_pts} 条数据 → {method_name} (末端月份={last_month})"
        )

        avg_fc, avg_ci, avg_rate = predict_fn(
            avg_clean, prediction_days, t_days_clean, last_month
        )

        max_available = is_metric_available(max_clean)
        if max_available:
            max_fc, max_ci, max_rate = predict_fn(
                max_clean, prediction_days, t_days_clean, last_month
            )
        else:
            max_fc = np.full(prediction_days, np.nan)
            max_ci, max_rate = 0.0, 0.0

        min_available = is_metric_available(min_clean)
        if min_available:
            min_fc, min_ci, min_rate = predict_fn(
                min_clean, prediction_days, t_days_clean, last_month
            )
        else:
            min_fc = np.full(prediction_days, np.nan)
            min_ci, min_rate = 0.0, 0.0

        # 4.5 环境因素分析（气温→粮温的滞后关联）
        outdoor_raw = [
            s.outdoor_temp
            for s in series_sorted
            if s.outdoor_temp is not None and s.outdoor_temp != 0.0
        ]
        env_context: Dict[str, Any] = {}
        env_trend_adjustment = 0.0

        if len(outdoor_raw) >= 3:
            outdoor_temps = np.array(outdoor_raw, dtype=float)
            # 过滤掉传感器故障值（物理边界）
            ot_valid = outdoor_temps[(outdoor_temps > -40) & (outdoor_temps < 60)]

            if len(ot_valid) >= 3 and float(np.std(ot_valid)) > 0.01:
                # 气温趋势斜率
                ot_t = np.arange(len(ot_valid), dtype=float)
                ot_coeffs = np.polyfit(ot_t, ot_valid, deg=1)
                ot_slope = float(ot_coeffs[0])  # °C/检测周期

                # 粮温趋势斜率
                avg_t = np.arange(len(avg_clean), dtype=float)
                avg_coeffs = np.polyfit(avg_t, avg_clean, deg=1)
                avg_slope = float(avg_coeffs[0])

                # 气温升速 - 粮温升速: 正值表示气温升得更快（粮温可能滞后跟涨）
                divergence = ot_slope - avg_slope

                if divergence > 0.1:
                    # 气温在升而粮温还没跟上 → 粮温可能滞后上升
                    env_trend_adjustment = min(divergence * 0.3, 0.5)
                    env_context["warning"] = (
                        f"气温上升速率 ({ot_slope:+.3f}°C/周期) 快于粮温 "
                        f"({avg_slope:+.3f}°C/周期)，粮温可能存在滞后升温风险"
                    )
                elif divergence < -0.1:
                    # 气温在降而粮温还没跟上 → 有利于后续降温
                    env_trend_adjustment = max(divergence * 0.3, -0.5)
                    env_context["info"] = (
                        f"气温下降速率快于粮温变化，有利于后续自然降温"
                    )

                env_context["outdoor_temp_trend"] = round(ot_slope, 4)
                env_context["grain_temp_trend"] = round(avg_slope, 4)
                env_context["trend_divergence"] = round(divergence, 4)
                env_context["latest_outdoor_temp"] = round(float(ot_valid[-1]), 1)

        # 当气温数据不可用时标注
        if not env_context:
            env_context["status"] = "unavailable"
            env_context["note"] = "气温数据缺失或全为零，无法进行环境关联分析"

        # 季节转折点检测 — 训练末端处于季节交接期时，预测方向可能反转
        # 秋→冬（10-11月）/ 冬→春（3-4月）是最容易出现方向性误判的时段
        transition_map = {
            10: "秋冬",
            11: "秋冬",
            3: "冬春",
            4: "冬春",
        }
        if last_month in transition_map:
            season_pair = transition_map[last_month]
            env_context["seasonal_transition"] = (
                f"当前处于{season_pair}季节转折期（{last_month}月），"
                f"粮温走势可能发生方向反转，统计外推存在较高不确定性，"
                f"建议结合现场巡检综合判断"
            )

        # 对 avg 预测施加环境修正（渐进式，第1天修正量小，最后一天最大）
        if env_trend_adjustment != 0:
            correction = np.linspace(
                0, env_trend_adjustment * prediction_days, prediction_days
            )
            avg_fc = avg_fc + correction
            # 重新计算修正后的日均变化率
            avg_rate = (
                float((avg_fc[-1] - avg_clean[-1]) / prediction_days)
                if prediction_days > 0
                else 0.0
            )

        # 5. 趋势判定（基于 avg 日均变化率，细化为 5 级）
        if avg_rate > 0.2:
            trend = "快速上升"
        elif avg_rate > 0.05:
            trend = "缓慢上升"
        elif avg_rate < -0.2:
            trend = "快速下降"
        elif avg_rate < -0.05:
            trend = "缓慢下降"
        else:
            trend = "稳定"

        # 6. 多因素风险评估
        mean_val = float(np.mean(avg_clean))
        cv = float(np.std(avg_clean, ddof=1) / mean_val) if mean_val != 0 else 0.0
        pred_max_val = round(float(max_fc[-1]), 1) if max_available else None
        risk = assess_risk(
            pred_avg=round(float(avg_fc[-1]), 1),
            pred_max=pred_max_val if pred_max_val is not None else 0.0,
            daily_rate=avg_rate,
            cv=cv,
            ci_half=avg_ci,
            horizon=prediction_days,
            n_pts=effective_pts,
            env_ctx=env_context if env_context.get("status") != "unavailable" else None,
        )

        # 7. 构建返回值
        # P1: 物理下界截断
        # 中国粮库的粮温在正常储粮条件下不会低于 -10°C（即便东北严寒地区也极少低于此值）
        # 如果预测值低于物理下界，说明算法在季节转折处做了不合理的外推，强制截断并标注
        GRAIN_TEMP_PHYSICAL_MIN = -10.0
        physical_clamp_applied = False
        for i in range(len(avg_fc)):
            if avg_fc[i] < GRAIN_TEMP_PHYSICAL_MIN:
                avg_fc[i] = GRAIN_TEMP_PHYSICAL_MIN
                physical_clamp_applied = True
        if physical_clamp_applied:
            # 重新计算修正后的日均变化率和趋势
            avg_rate = (
                float((avg_fc[-1] - avg_clean[-1]) / prediction_days)
                if prediction_days > 0
                else 0.0
            )
            if avg_rate > 0.2:
                trend = "快速上升"
            elif avg_rate > 0.05:
                trend = "缓慢上升"
            elif avg_rate < -0.2:
                trend = "快速下降"
            elif avg_rate < -0.05:
                trend = "缓慢下降"
            else:
                trend = "稳定"
            logger.warning(
                f"预测值低于物理下界 {GRAIN_TEMP_PHYSICAL_MIN}°C，已截断至下界。"
                f"这通常发生在季节转折点的过度外推，请结合现场判断。"
            )

        predicted_avg = round(float(avg_fc[-1]), 1)

        daily_preds = [
            {
                "day": d + 1,
                "avg_temp": round(float(avg_fc[d]), 1),
                "max_temp": round(float(max_fc[d]), 1) if max_available else None,
                "min_temp": round(float(min_fc[d]), 1) if min_available else None,
            }
            for d in range(prediction_days)
        ]

        # 兼容旧消费方的 risk_assessment 一句话
        risk_level = risk["risk_level"]
        if risk_level == "critical":
            risk_text = "高风险，建议立即采取降温措施"
        elif risk_level == "high":
            risk_text = "持续升温风险，建议加强监控"
        elif risk_level == "medium":
            risk_text = "存在一定风险，建议持续关注"
        else:
            risk_text = "状态良好"

        # P1: 空间热点分析（仅在显式请求时执行）
        spatial_result = None
        if include_spatial:
            try:
                from app.services.predictor.spatial_predictor import (
                    SpatialTempPredictor,
                )

                spatial_predictor = SpatialTempPredictor()
                sensors = spatial_predictor.parse_temp_values(series_sorted)

                if len(sensors) >= 3:
                    hotspots = spatial_predictor.identify_hotspots(sensors)
                    spatial_preds = spatial_predictor.predict_spatial(
                        sensors, prediction_days
                    )

                    spatial_result = {
                        "sensor_count": len(sensors),
                        "hotspots": hotspots,
                        "top_5_predictions": spatial_preds[
                            :5
                        ],  # 只返回前 5 个最高温点位
                    }
                else:
                    spatial_result = {
                        "status": "insufficient_sensors",
                        "message": f"仅检测到 {len(sensors)} 个传感器（< 3），无法进行空间分析",
                    }
            except Exception as e:
                logger.warning(f"空间分析失败: {e}")
                spatial_result = {"status": "error", "message": str(e)}

        return {
            # ── 保留字段 (report() / LLM 旧调用方直接使用) ──
            "silo_id": silo_id,
            "prediction_days": prediction_days,
            "current_avg_temp": round(float(avg_clean[-1]), 1),
            "predicted_avg_temp": predicted_avg,
            "trend": trend,
            "risk_assessment": risk_text,
            "context": {
                "silo_id": silo_id,
                "current_temp": round(float(avg_clean[-1]), 1),
                "historical_trend": trend,
                "daily_rate": round(avg_rate, 4),
                "predicted_temp_in_days": predicted_avg,
                "prediction_period": f"{prediction_days}天",
            },
            # ── 新增字段 ──
            "method": method_name,
            "confidence_interval": {
                "lower": round(predicted_avg - avg_ci, 1),
                "upper": round(predicted_avg + avg_ci, 1),
                "level": 0.95,
            },
            "daily_predictions": daily_preds,
            "multi_metric": {
                "avg_temp": {
                    "current": round(float(avg_clean[-1]), 1),
                    "predicted": predicted_avg,
                    "daily_rate": round(avg_rate, 4),
                },
                "max_temp": {
                    "current": round(float(max_clean[-1]), 1)
                    if max_available
                    else None,
                    "predicted": round(float(max_fc[-1]), 1) if max_available else None,
                    "daily_rate": round(max_rate, 4) if max_available else None,
                    "status": "available"
                    if max_available
                    else "数据不可用(API未返回有效最高粮温)",
                },
                "min_temp": {
                    "current": round(float(min_clean[-1]), 1)
                    if min_available
                    else None,
                    "predicted": round(float(min_fc[-1]), 1) if min_available else None,
                    "daily_rate": round(min_rate, 4) if min_available else None,
                    "status": "available"
                    if min_available
                    else "数据不可用(API未返回有效最低粮温)",
                },
            },
            "data_quality": {
                "total_points": total_pts,
                "outliers_removed": outliers_cnt,
                "effective_points": effective_pts,
                "experimental": is_experimental,  # P0-2: 试验性标记
                "anomaly_details": anomaly_stats,  # P0-1: 异常检测统计
                "note": _build_data_quality_note(
                    effective_pts, physical_clamp_applied, is_experimental
                ),
            },
            "environmental_context": env_context,
            "risk_details": risk,
            "spatial_analysis": spatial_result,  # P1: 空间热点分析（可选）
        }
    

    # T12: LLM 辅助温度预测
    def llm_temperature_prediction(self, silo_id: str, prediction_days: int = 5,
                                    start_time: Optional[str] = None, end_time: Optional[str] = None) -> Dict[str, Any]:
        """
        T12: LLM 辅助温度预测

        将历史粮温数据交给 LLM 进行智能分析，生成逐日温度预测、趋势判断、
        风险评估和储粮建议，并输出包含历史+预测双段曲线的可视化图表。
        """
        house_code = self.wms_client.resolve_house_code(silo_id)

        # 确定历史参考时间范围
        if start_time and end_time:
            ref_start = self._parse_dt(start_time)
            ref_end = self._parse_dt(end_time)
        else:
            ref_end = datetime.now()
            ref_start = ref_end - timedelta(days=14)

        series = self.wms_client.get_grain_temperature(house_code, ref_start, ref_end)

        # 数据不足熔断
        if not series or len(series) <= 2:
            msg = (f"由于历史检测记录仅有 {len(series) if series else 0} 条（不足3条），"
                   "暂无法为您提供可靠的 LLM 辅助温度预测，请在积累更多数据后再试。")
            logger.info(msg)
            return {"silo_id": silo_id, "status": "no_data", "message": msg}

        # 确保按时间排序
        series = sorted(series, key=lambda x: x.check_time)

        prediction_days = max(3, min(7, prediction_days))

        # 构造历史数据序列（供 LLM 分析）
        history_records = []
        for s in series:
            history_records.append({
                "check_time": s.check_time,
                "avg_temp": round(s.avg_temp, 2),
                "max_temp": round(s.max_temp, 2),
                "min_temp": round(s.min_temp, 2),
                "indoor_temp": round(s.indoor_temp, 2) if s.indoor_temp else None,
                "outdoor_temp": round(s.outdoor_temp, 2) if s.outdoor_temp else None,
            })

        context = {
            "silo_id": silo_id,
            "prediction_days": prediction_days,
            "history_count": len(history_records),
            "history_data": history_records,
        }

        prompt_template = (
            f"你是一位粮食储藏温度预测专家。以下是仓号 {silo_id} 的历史粮温数据（共 {len(history_records)} 条），"
            f"请根据这些数据预测未来 {prediction_days} 天的逐日平均粮温。\n\n"
            "请严格以如下 JSON 格式返回结果（不要包含任何其他内容）：\n"
            "{\n"
            '  "trend": "上升" | "下降" | "稳定",\n'
            '  "trend_analysis": "对趋势的简要分析说明",\n'
            '  "predicted_temps": [{"day": 1, "avg_temp": 25.3}, ...],\n'
            '  "risk_level": "low" | "medium" | "high",\n'
            '  "risk_reason": "风险原因说明",\n'
            '  "recommendations": ["建议1", "建议2", ...],\n'
            '  "confidence": "high" | "medium" | "low"\n'
            "}\n\n"
            "注意：\n"
            "- predicted_temps 数组长度必须等于预测天数\n"
            "- avg_temp 为预测的平均粮温（℃），保留1位小数\n"
            "- 基于温度变化速率、季节性规律和粮食储藏特点做出合理预测\n"
            "- 如果数据量较少或波动较大，请降低 confidence"
        )

        # 调用 LLM 进行推理
        llm_result = self.llm_service.reason_with_context(context, prompt_template)

        # 解析 LLM 返回结果
        predicted_temps = []
        trend = "稳定"
        trend_analysis = ""
        risk_level = "low"
        risk_reason = ""
        recommendations = []
        confidence = "medium"
        llm_success = False

        try:
            # reason_with_context 返回的可能是已经解析的 dict 或含 conclusion 的 dict
            result_data = llm_result
            if isinstance(result_data, dict):
                # 如果 LLM 返回了 conclusion 字段（字符串），尝试解析它
                if "conclusion" in result_data and isinstance(result_data["conclusion"], str):
                    try:
                        conclusion_str = result_data["conclusion"]
                        if "```json" in conclusion_str:
                            conclusion_str = conclusion_str.split("```json")[1].split("```")[0]
                        elif "```" in conclusion_str:
                            conclusion_str = conclusion_str.split("```")[1].split("```")[0]
                        result_data = json.loads(conclusion_str)
                    except (json.JSONDecodeError, IndexError):
                        pass

                if "predicted_temps" in result_data:
                    raw_preds = result_data["predicted_temps"]
                    if isinstance(raw_preds, list) and len(raw_preds) > 0:
                        for p in raw_preds:
                            if isinstance(p, dict) and "avg_temp" in p:
                                predicted_temps.append({
                                    "day": p.get("day", len(predicted_temps) + 1),
                                    "avg_temp": round(float(p["avg_temp"]), 1),
                                })
                        if predicted_temps:
                            llm_success = True

                trend = result_data.get("trend", trend)
                trend_analysis = result_data.get("trend_analysis", trend_analysis)
                risk_level = result_data.get("risk_level", risk_level)
                risk_reason = result_data.get("risk_reason", risk_reason)
                recommendations = result_data.get("recommendations", recommendations)
                confidence = result_data.get("confidence", confidence)
        except Exception as e:
            logger.warning(f"解析 LLM 预测结果失败: {e}")

        # LLM 返回异常时 fallback 到线性外推
        if not llm_success or len(predicted_temps) < prediction_days:
            logger.info("LLM 预测结果不完整，回退到线性外推")
            temps_list = [s.avg_temp for s in series]
            overall_diff = temps_list[-1] - temps_list[0]
            daily_rate = overall_diff / max(len(temps_list) - 1, 1)
            predicted_temps = []
            for d in range(1, prediction_days + 1):
                predicted_temps.append({
                    "day": d,
                    "avg_temp": round(temps_list[-1] + daily_rate * d, 1),
                })
            if not trend_analysis:
                trend = "上升" if daily_rate > 0.05 else ("下降" if daily_rate < -0.05 else "稳定")
                trend_analysis = f"基于线性外推（日均变化 {daily_rate:+.2f}°C）"
            confidence = "low"
            if not risk_reason:
                risk_level = "medium" if abs(daily_rate) > 0.2 else "low"
                risk_reason = "线性外推结果，仅供参考" if not risk_reason else risk_reason
            if not recommendations:
                recommendations = ["建议持续关注温度变化趋势", "如温度持续升高请及时通风降温"]

        # 生成图表
        chart_result = self._generate_prediction_chart(
            silo_id=silo_id,
            history_series=series,
            predicted_temps=predicted_temps,
            risk_level=risk_level,
        )

        return {
            "silo_id": silo_id,
            "status": "success",
            "prediction_days": prediction_days,
            "current_avg_temp": round(series[-1].avg_temp, 1),
            "trend": trend,
            "trend_analysis": trend_analysis,
            "predicted_temps": predicted_temps,
            "risk_level": risk_level,
            "risk_reason": risk_reason,
            "recommendations": recommendations,
            "confidence": confidence,
            "history_count": len(series),
            "chart_path": chart_result.get("file_path"),
            "chart_status": chart_result.get("status"),
            "method": "llm" if llm_success else "linear_fallback",
        }

    def _generate_prediction_chart(self, silo_id: str, history_series: list,
                                    predicted_temps: List[Dict[str, Any]],
                                    risk_level: str = "low") -> Dict[str, Any]:
        """生成历史+预测双段曲线图表。"""
        plt.rcParams['font.sans-serif'] = ['SimHei', 'Microsoft YaHei', 'SimSun', 'STHeiti', 'Arial Unicode MS', 'DejaVu Sans']
        plt.rcParams['axes.unicode_minus'] = False

        out_dir = Path(__file__).resolve().parents[2] / "artifacts" / "charts"
        out_dir.mkdir(parents=True, exist_ok=True)
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        out_path = out_dir / f"{silo_id}_llm_prediction_{ts}.png"

        try:
            # --- 历史数据段 ---
            hist_times = []
            hist_avg = []
            hist_max = []
            hist_min = []
            for record in history_series:
                try:
                    ct = self._parse_dt(record.check_time)
                    hist_times.append(ct)
                    hist_avg.append(record.avg_temp)
                    hist_max.append(record.max_temp)
                    hist_min.append(record.min_temp)
                except Exception:
                    continue

            if not hist_times:
                return {"file_path": str(out_path), "status": "no_data"}

            # 限制历史段最多 7 个点
            if len(hist_times) > 7:
                hist_times = hist_times[-7:]
                hist_avg = hist_avg[-7:]
                hist_max = hist_max[-7:]
                hist_min = hist_min[-7:]

            # --- 预测数据段 ---
            last_hist_time = hist_times[-1]
            pred_times = []
            pred_avg = []
            for p in predicted_temps:
                pred_times.append(last_hist_time + timedelta(days=p["day"]))
                pred_avg.append(p["avg_temp"])

            # 风险颜色
            risk_colors = {"low": "#00AA00", "medium": "#FF8C00", "high": "#CC0000"}
            pred_color = risk_colors.get(risk_level, "#FF8C00")

            # --- 绘图 ---
            fig, ax = plt.subplots(figsize=(10, 6))

            # 历史段：实线 + 温度范围填充带
            ax.plot(hist_times, hist_avg, color="#0066CC", linewidth=2.0, marker="o",
                    markersize=4, label="历史平均粮温", zorder=3)
            ax.fill_between(hist_times, hist_min, hist_max, color="#0066CC", alpha=0.15,
                            label="历史温度范围（最低~最高）")

            # 连接点（历史末尾到预测起点）
            bridge_times = [hist_times[-1], pred_times[0]]
            bridge_avg = [hist_avg[-1], pred_avg[0]]
            ax.plot(bridge_times, bridge_avg, color=pred_color, linewidth=1.5,
                    linestyle="--", alpha=0.6, zorder=2)

            # 预测段：虚线
            ax.plot(pred_times, pred_avg, color=pred_color, linewidth=2.0, marker="D",
                    markersize=5, linestyle="--", label=f"LLM预测粮温（风险: {risk_level}）", zorder=3)

            # 垂直分界线
            ax.axvline(x=hist_times[-1], color="gray", linewidth=1.0, linestyle=":",
                       alpha=0.7, zorder=1)
            y_mid = (ax.get_ylim()[0] + ax.get_ylim()[1]) / 2
            ax.text(hist_times[-1], ax.get_ylim()[1] * 0.98, " 历史 | 预测 ",
                    ha="center", va="top", fontsize=9, color="gray",
                    bbox=dict(boxstyle="round,pad=0.3", facecolor="white", edgecolor="gray", alpha=0.8))

            # 坐标轴与标题
            ax.set_xlabel("日期", fontsize=12)
            ax.set_ylabel("温度 (℃)", fontsize=12)
            ax.set_title(f"{silo_id}号仓 - LLM 辅助温度预测", fontsize=14, fontweight="bold")

            # X 轴格式
            all_times = hist_times + pred_times
            if len(all_times) <= 14:
                ax.set_xticks(all_times)
                ax.set_xticklabels([t.strftime("%m-%d") for t in all_times], rotation=45, ha="right")
            else:
                ax.xaxis.set_major_formatter(mdates.DateFormatter("%m-%d"))
                plt.xticks(rotation=45, ha="right")

            ax.legend(loc="upper left", fontsize=9, frameon=True, fancybox=True, shadow=True)
            ax.grid(True, linestyle="--", alpha=0.5)
            plt.tight_layout()
            plt.savefig(out_path, dpi=150, bbox_inches="tight")
            plt.close()

            logger.info(f"LLM 预测图表生成成功: {out_path}")
            return {"file_path": str(out_path), "status": "generated"}

        except Exception as e:
            logger.error(f"生成 LLM 预测图表失败: {e}", exc_info=True)
            plt.close("all")
            return {"file_path": str(out_path), "status": "error", "message": str(e)}

    # V008: 三温图生成（气温、仓温、粮温）
    def generate_three_temp_chart(self, silo_id: str, time_range_hours: Optional[int] = None, start_time: Optional[datetime] = None, end_time: Optional[datetime] = None) -> Dict[str, Any]:
        """
        V008: 生成三温图
        """
        # 配置中文字体，确保图表文字显示正常
        plt.rcParams['font.sans-serif'] = ['SimHei', 'Microsoft YaHei', 'SimSun', 'STHeiti', 'Arial Unicode MS', 'DejaVu Sans']
        plt.rcParams['axes.unicode_minus'] = False

        house_code = self.wms_client.resolve_house_code(silo_id)
        
        # ... (确定时间范围逻辑保持不变)
        if start_time is not None and end_time is not None:
            query_start_time = start_time
            query_end_time = end_time
        elif time_range_hours is not None:
            query_end_time = datetime.now()
            query_start_time = query_end_time - timedelta(hours=time_range_hours)
        else:
            query_end_time = datetime.now()
            query_start_time = query_end_time - timedelta(hours=24)
        
        series = self.wms_client.get_grain_temperature(house_code, query_start_time, query_end_time)
        
        # V008: 落实规范 - 数据不足熔断 (<= 1)
        if not series or len(series) <= 1:
            msg = f"当前该仓房({silo_id})仅有 {len(series) if series else 0} 次检测记录，无法绘制多点对比趋势图，请检查数据采集状态。"
            logger.info(msg)
            return {"silo_id": silo_id, "chart_type": "three_temp", "status": "no_data", "message": msg}

        out_dir = Path(__file__).resolve().parents[2] / "artifacts" / "charts"
        out_dir.mkdir(parents=True, exist_ok=True)
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        out_path = out_dir / f"{silo_id}_three_temp_{ts}.png"
        
        if not series or len(series) < 2:
            logger.warning(f"数据不足，无法生成三温图（仓号: {silo_id}）")
            return {"silo_id": silo_id, "chart_type": "three_temp", "file_path": str(out_path), "status": "no_data"}
        
        try:
            # 解析时间并提取数据
            times = []
            outdoor_temps = []
            indoor_temps = []
            avg_temps = []
            
            for record in series:
                try:
                    check_time = datetime.strptime(record.check_time, "%Y-%m-%d %H:%M:%S")
                    times.append(check_time)
                    outdoor_temps.append(record.outdoor_temp)
                    indoor_temps.append(record.indoor_temp)
                    avg_temps.append(record.avg_temp)
                except (ValueError, AttributeError) as e:
                    logger.warning(f"解析数据记录失败: {e}")
                    continue
            
            if len(times) < 2:
                return {"silo_id": silo_id, "chart_type": "three_temp", "file_path": str(out_path), "status": "no_data"}
            
            # V008: 限制为最多7次数据（按规范要求）
            # 首先确保数据按时间排序（从早到晚）
            sorted_data = sorted(zip(times, outdoor_temps, indoor_temps, avg_temps), key=lambda x: x[0])
            times, outdoor_temps, indoor_temps, avg_temps = zip(*sorted_data)
            times = list(times)
            outdoor_temps = list(outdoor_temps)
            indoor_temps = list(indoor_temps)
            avg_temps = list(avg_temps)
            
            # 如果超过7个数据点，根据是否指定了时间范围来决定取哪些数据
            if len(times) > 7:
                if start_time is not None and end_time is not None:
                    # 用户指定了时间范围：从最早的数据开始，取前7个
                    # 这样可以显示用户指定时间范围内的早期数据趋势
                    times = times[:7]
                    outdoor_temps = outdoor_temps[:7]
                    indoor_temps = indoor_temps[:7]
                    avg_temps = avg_temps[:7]
                    logger.info(f"数据点超过7个，已限制为最早的7次数据点（从 {len(sorted_data)} 个减少到 7 个，时间范围: {start_time.strftime('%Y-%m-%d')} 到 {end_time.strftime('%Y-%m-%d')}）")
                else:
                    # 未指定时间范围（使用 time_range_hours）：取最近的7个（时间最晚的7个）
                    times = times[-7:]
                    outdoor_temps = outdoor_temps[-7:]
                    indoor_temps = indoor_temps[-7:]
                    avg_temps = avg_temps[-7:]
                    logger.info(f"数据点超过7个，已限制为最近的7次数据点（从 {len(sorted_data)} 个减少到 7 个）")
            
            # 创建图表
            fig, ax = plt.subplots(figsize=(10, 6))
            
            # 绘制三条曲线，使用不同颜色
            ax.plot(times, outdoor_temps, label="气温", color="#0066CC", linewidth=2.0, marker="o", markersize=4)
            ax.plot(times, indoor_temps, label="仓温", color="#00AA00", linewidth=2.0, marker="s", markersize=4)
            ax.plot(times, avg_temps, label="粮温", color="#CC0000", linewidth=2.0, marker="^", markersize=4)
            
            # 设置坐标轴
            ax.set_xlabel("检测日期", fontsize=12)
            ax.set_ylabel("温度 (℃)", fontsize=12)
            ax.set_title(f"{silo_id}号仓 - 三温图", fontsize=14, fontweight="bold")
            
            # 格式化 X 轴时间 - 根据数据点数量和时间跨度自适应
            # 如果数据点较少，显示所有日期；如果较多，自动调整间隔
            if len(times) <= 7:
                # 数据点少，显示所有日期
                ax.set_xticks(times)
                ax.set_xticklabels([t.strftime("%Y-%m-%d") for t in times], rotation=45, ha="right")
            else:
                # 数据点多，使用自动定位器
                ax.xaxis.set_major_formatter(mdates.DateFormatter("%Y-%m-%d"))
                # 根据时间跨度选择合适的间隔
                time_span = (max(times) - min(times)).days
                if time_span <= 7:
                    interval = 1  # 每天
                elif time_span <= 30:
                    interval = max(1, time_span // 7)  # 每周
                else:
                    interval = max(1, time_span // 10)  # 每月
                ax.xaxis.set_major_locator(mdates.DayLocator(interval=interval))
                plt.xticks(rotation=45, ha="right")
            
            # 添加图例（右上角）
            ax.legend(loc="upper right", fontsize=10, frameon=True, fancybox=True, shadow=True)
            
            # 添加网格线
            ax.grid(True, linestyle="--", alpha=0.5)
            
            # 调整布局
            plt.tight_layout()
            
            # 保存图表
            plt.savefig(out_path, dpi=150, bbox_inches="tight")
            plt.close()
            
            logger.info(f"三温图生成成功: {out_path}")
            return {"silo_id": silo_id, "chart_type": "three_temp", "file_path": str(out_path), "status": "generated"}
            
        except Exception as e:
            logger.error(f"生成三温图失败: {e}", exc_info=True)
            return {"silo_id": silo_id, "chart_type": "three_temp", "file_path": str(out_path), "status": "error"}

    # V008: 两湿图生成（气湿、仓湿）
    def generate_two_humidity_chart(self, silo_id: str, time_range_hours: Optional[int] = None, start_time: Optional[datetime] = None, end_time: Optional[datetime] = None) -> Dict[str, Any]:
        """
        V008: 生成两湿图
        """
        house_code = self.wms_client.resolve_house_code(silo_id)
        
        # 确定时间范围
        if start_time is not None and end_time is not None:
            query_start_time = start_time
            query_end_time = end_time
        elif time_range_hours is not None:
            query_end_time = datetime.now()
            query_start_time = query_end_time - timedelta(hours=time_range_hours)
        else:
            query_end_time = datetime.now()
            query_start_time = query_end_time - timedelta(hours=24)
        
        series = self.wms_client.get_grain_temperature(house_code, query_start_time, query_end_time)
        
        # V008: 落实规范 - 数据不足熔断 (<= 1)
        if not series or len(series) <= 1:
            msg = f"当前该仓房({silo_id})仅有 {len(series) if series else 0} 次检测记录，无法绘制两湿图，请检查数据采集状态。"
            logger.info(msg)
            return {"silo_id": silo_id, "chart_type": "two_humidity", "status": "no_data", "message": msg}

        # 配置中文字体，增加更多 Windows 常见中文字体候选项
        plt.rcParams['font.sans-serif'] = ['SimHei', 'Microsoft YaHei', 'SimSun', 'STHeiti', 'Arial Unicode MS', 'DejaVu Sans']
        plt.rcParams['axes.unicode_minus'] = False  # 解决负号显示问题

        logger.info(f"两湿图查询到 {len(series)} 条记录，时间范围: {query_start_time.strftime('%Y-%m-%d')} 到 {query_end_time.strftime('%Y-%m-%d')}")
        
        out_dir = Path(__file__).resolve().parents[2] / "artifacts" / "charts"
        out_dir.mkdir(parents=True, exist_ok=True)
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        out_path = out_dir / f"{silo_id}_two_humidity_{ts}.png"
        
        if not series or len(series) < 2:
            logger.warning(f"数据不足，无法生成两湿图（仓号: {silo_id}）")
            return {"silo_id": silo_id, "chart_type": "two_humidity", "file_path": str(out_path), "status": "no_data"}
        
        try:
            # 解析时间并提取数据
            times = []
            outdoor_humidities = []
            indoor_humidities = []
            
            for record in series:
                try:
                    check_time = datetime.strptime(record.check_time, "%Y-%m-%d %H:%M:%S")
                    times.append(check_time)
                    outdoor_humidities.append(record.outdoor_humidity)
                    indoor_humidities.append(record.indoor_humidity)
                except (ValueError, AttributeError) as e:
                    logger.warning(f"解析数据记录失败: {e}")
                    continue
            
            if len(times) < 2:
                return {"silo_id": silo_id, "chart_type": "two_humidity", "file_path": str(out_path), "status": "no_data"}
            
            # V008: 限制为最多7次数据（按规范要求）
            # 首先确保数据按时间排序（从早到晚）
            sorted_data = sorted(zip(times, outdoor_humidities, indoor_humidities), key=lambda x: x[0])
            times, outdoor_humidities, indoor_humidities = zip(*sorted_data)
            times = list(times)
            outdoor_humidities = list(outdoor_humidities)
            indoor_humidities = list(indoor_humidities)
            
            # 如果超过7个数据点，根据是否指定了时间范围来决定取哪些数据
            if len(times) > 7:
                if start_time is not None and end_time is not None:
                    # 用户指定了时间范围：从最早的数据开始，取前7个
                    # 这样可以显示用户指定时间范围内的早期数据趋势
                    logger.info(f"两湿图：用户指定了时间范围，数据点超过7个（共{len(times)}个），取最早的7个数据点")
                    logger.info(f"  最早数据点: {times[0].strftime('%Y-%m-%d %H:%M:%S')}, 最晚数据点: {times[-1].strftime('%Y-%m-%d %H:%M:%S')}")
                    times = times[:7]
                    outdoor_humidities = outdoor_humidities[:7]
                    indoor_humidities = indoor_humidities[:7]
                    logger.info(f"  取前7个数据点: {[t.strftime('%Y-%m-%d') for t in times]}")
                else:
                    # 未指定时间范围（使用 time_range_hours）：取最近的7个（时间最晚的7个）
                    logger.info(f"两湿图：未指定时间范围，数据点超过7个（共{len(times)}个），取最近的7个数据点")
                    times = times[-7:]
                    outdoor_humidities = outdoor_humidities[-7:]
                    indoor_humidities = indoor_humidities[-7:]
                    logger.info(f"  取后7个数据点: {[t.strftime('%Y-%m-%d') for t in times]}")
            
            # 创建图表
            fig, ax = plt.subplots(figsize=(10, 6))
            
            # 绘制两条曲线，使用不同颜色
            ax.plot(times, outdoor_humidities, label="气湿", color="#0066CC", linewidth=2.0, marker="o", markersize=4)
            ax.plot(times, indoor_humidities, label="仓湿", color="#00AA00", linewidth=2.0, marker="s", markersize=4)
            
            # 设置坐标轴
            ax.set_xlabel("检测日期", fontsize=12)
            ax.set_ylabel("湿度 (%)", fontsize=12)
            ax.set_title(f"{silo_id}号仓 - 两湿图", fontsize=14, fontweight="bold")
            
            # Y 轴范围：0-100%
            ax.set_ylim(0, 100)
            
            # 格式化 X 轴时间 - 根据数据点数量和时间跨度自适应
            # 如果数据点较少，显示所有日期；如果较多，自动调整间隔
            if len(times) <= 7:
                # 数据点少，显示所有日期
                ax.set_xticks(times)
                ax.set_xticklabels([t.strftime("%Y-%m-%d") for t in times], rotation=45, ha="right")
            else:
                # 数据点多，使用自动定位器
                ax.xaxis.set_major_formatter(mdates.DateFormatter("%Y-%m-%d"))
                # 根据时间跨度选择合适的间隔
                time_span = (max(times) - min(times)).days
                if time_span <= 7:
                    interval = 1  # 每天
                elif time_span <= 30:
                    interval = max(1, time_span // 7)  # 每周
                else:
                    interval = max(1, time_span // 10)  # 每月
                ax.xaxis.set_major_locator(mdates.DayLocator(interval=interval))
                plt.xticks(rotation=45, ha="right")
            
            # 添加图例（右上角）
            ax.legend(loc="upper right", fontsize=10, frameon=True, fancybox=True, shadow=True)
            
            # 添加网格线
            ax.grid(True, linestyle="--", alpha=0.5)
            
            # 调整布局
            plt.tight_layout()
            
            # 保存图表
            plt.savefig(out_path, dpi=150, bbox_inches="tight")
            plt.close()
            
            logger.info(f"两湿图生成成功: {out_path}")
            return {"silo_id": silo_id, "chart_type": "two_humidity", "file_path": str(out_path), "status": "generated"}
            
        except Exception as e:
            logger.error(f"生成两湿图失败: {e}", exc_info=True)
            return {"silo_id": silo_id, "chart_type": "two_humidity", "file_path": str(out_path), "status": "error"}

    # T8: 报告生成（生成 docx 到本地 artifacts/reports）
    def report(self, silo_ids: List[str], report_type: str = "daily", start_time: Optional[str] = None, end_time: Optional[str] = None) -> Dict[str, Any]:
        """
        T8: 报告生成

        生成粮情分析报告，包含完整分析文字和图表

        Args:
            silo_ids: 仓号列表
            report_type: 报告类型（daily/weekly/alert）
            start_time: 开始时间字符串，格式：'YYYY-MM-DD HH:MM:SS' 或 'YYYY-MM-DD'（可选）
            end_time: 结束时间字符串，格式：'YYYY-MM-DD HH:MM:SS' 或 'YYYY-MM-DD'（可选）

        Returns:
            报告下载链接
        """
        logger.info(f"report method called: silo_ids={silo_ids}, report_type={report_type}")
        from docx import Document  # type: ignore
        from docx.shared import Inches  # type: ignore

        out_dir = Path(__file__).resolve().parents[2] / "artifacts" / "reports"
        out_dir.mkdir(parents=True, exist_ok=True)
        report_id = f"{'_'.join(silo_ids)}_{report_type}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        out_path = out_dir / f"{report_id}.docx"

        doc = Document()
        doc.add_heading(f"粮情分析报告 ({report_type})", level=0)
        doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

        for silo_id in silo_ids:
            doc.add_heading(f"{silo_id}号仓", level=1)
            try:
                info = self.wms_client.get_warehouse_info(silo_id)
                doc.add_paragraph(f"仓房：{info.house_name} / 库点：{info.depot_name}")
            except Exception:
                pass

            # 解析时间范围
            if start_time and end_time:
                # 使用用户指定的时间范围
                try:
                    report_start_time = self._parse_dt(start_time)
                    report_end_time = self._parse_dt(end_time)
                    # 如果只提供了日期，设置时间部分
                    if len(start_time) == 10:  # YYYY-MM-DD
                        report_start_time = report_start_time.replace(hour=0, minute=0, second=0)
                    if len(end_time) == 10:  # YYYY-MM-DD
                        report_end_time = report_end_time.replace(hour=23, minute=59, second=59)
                except Exception as e:
                    logger.warning(f"解析时间范围失败: {e}，使用默认24小时")
                    report_end_time = datetime.now()
                    report_start_time = report_end_time - timedelta(hours=24)
            else:
                # 默认使用最近24小时
                report_end_time = datetime.now()
                report_start_time = report_end_time - timedelta(hours=24)
            
            # 粮温与环境摘要（使用指定的时间范围）
            temps = self.wms_client.get_grain_temperature(silo_id, report_start_time, report_end_time)
            if temps:
                # 取最近一次检测的环境数据
                latest = temps[-1]
                avg_avg = sum(t.avg_temp for t in temps) / len(temps)
                max_max = max(t.max_temp for t in temps)
                min_min = min(t.min_temp for t in temps)
                
                time_span_str = f"{report_start_time.strftime('%Y-%m-%d')} 到 {report_end_time.strftime('%Y-%m-%d')}"
                doc.add_paragraph(f"指定时间段（{time_span_str}）摘要：")
                doc.add_paragraph(f"  - 平均粮温：{avg_avg:.1f}°C (范围: {min_min:.1f}°C ~ {max_max:.1f}°C)")
                doc.add_paragraph(f"  - 仓内状态：温度 {latest.indoor_temp:.1f}°C，湿度 {latest.indoor_humidity:.1f}%")
                doc.add_paragraph(f"  - 仓外环境：温度 {latest.outdoor_temp:.1f}°C，湿度 {latest.outdoor_humidity:.1f}%")

            # V008: 生成三温图（使用指定的时间范围）
            try:
                logger.info(f"生成三温图，时间范围: {report_start_time.strftime('%Y-%m-%d %H:%M:%S')} 到 {report_end_time.strftime('%Y-%m-%d %H:%M:%S')}")
                three_temp_result = self.generate_three_temp_chart(silo_id, start_time=report_start_time, end_time=report_end_time)
                if three_temp_result.get("status") == "generated" and three_temp_result.get("file_path"):
                    chart_path = Path(three_temp_result["file_path"])
                    if chart_path.exists():
                        doc.add_paragraph("三温图：")
                        doc.add_picture(str(chart_path), width=Inches(6))
                    else:
                        logger.warning(f"三温图文件不存在: {chart_path}")
                elif three_temp_result.get("status") == "no_data":
                    doc.add_paragraph("三温图：数据不足，无法生成图表")
            except Exception as e:
                logger.warning(f"生成三温图失败: {e}")

            # V008: 生成两湿图（使用指定的时间范围）
            try:
                logger.info(f"生成两湿图，时间范围: {report_start_time.strftime('%Y-%m-%d %H:%M:%S')} 到 {report_end_time.strftime('%Y-%m-%d %H:%M:%S')}")
                two_humidity_result = self.generate_two_humidity_chart(silo_id, start_time=report_start_time, end_time=report_end_time)
                if two_humidity_result.get("status") == "generated" and two_humidity_result.get("file_path"):
                    chart_path = Path(two_humidity_result["file_path"])
                    if chart_path.exists():
                        doc.add_paragraph("两湿图：")
                        doc.add_picture(str(chart_path), width=Inches(6))
                    else:
                        logger.warning(f"两湿图文件不存在: {chart_path}")
                elif two_humidity_result.get("status") == "no_data":
                    doc.add_paragraph("两湿图：数据不足，无法生成图表")
            except Exception as e:
                logger.warning(f"生成两湿图失败: {e}")

            # 生成温度曲线图（使用指定的时间范围）
            try:
                logger.info(f"生成趋势图，时间范围: {report_start_time.strftime('%Y-%m-%d %H:%M:%S')} 到 {report_end_time.strftime('%Y-%m-%d %H:%M:%S')}")
                chart_result = self.visualization(silo_id, chart_type="line", start_time=report_start_time, end_time=report_end_time)
                if chart_result.get("status") == "generated" and chart_result.get("file_path"):
                    chart_path = Path(chart_result["file_path"])
                    if chart_path.exists():
                        doc.add_paragraph("温度变化趋势图：")
                        doc.add_picture(str(chart_path), width=Inches(6))
                    else:
                        logger.warning(f"Chart file not found: {chart_path}")
            except Exception as e:
                logger.warning(f"Failed to add chart for {silo_id}: {e}")

            # 引用分析工具结果
            try:
                # 自动提取指定时间段的数据
                extraction_result = self.extraction(silo_id, start_time=report_start_time, end_time=report_end_time)
                warehouse_info = extraction_result.get("warehouse_info", {})
                
                # 1. 仓房基本信息章节
                doc.add_heading("1. 仓房基本情况", level=2)
                doc.add_paragraph(f"仓房编号：{warehouse_info.get('house_code', silo_id)}")
                doc.add_paragraph(f"仓房名称：{warehouse_info.get('house_name', '未知')}")
                doc.add_paragraph(f"仓房类型：{warehouse_info.get('house_type_name', '未知')}")
                doc.add_paragraph(f"储粮性质：{warehouse_info.get('grain_nature', '未知')}")
                doc.add_paragraph(f"粮食品种：{warehouse_info.get('variety', '未知')}")
                doc.add_paragraph(f"设计仓容：{warehouse_info.get('design_capacity', 0)} 吨")
                doc.add_paragraph(f"实际粮高：{warehouse_info.get('actual_grain_height', 0)} 米")

                # 2. 详细粮情分析
                a = self.analysis(silo_id, start_time=report_start_time, end_time=report_end_time)
                doc.add_heading("2. 智能风险评估", level=2)
                
                # 检查数据量是否充足
                data_count = len(extraction_result.get("grain_temperature", []))
                if data_count < 7:
                    doc.add_paragraph(f"【重要提示】当前时间段内仅获取到 {data_count} 次检测数据（少于建议的 7 次），分析结论可能存在偏差，仅供参考。")
                
                doc.add_paragraph(f"风险等级：{a.get('risk_level','unknown')}")
                doc.add_paragraph(f"健康评分：{a.get('score','-')} / 100")
                
                findings = a.get("findings") or []
                if findings:
                    doc.add_paragraph("主要监测发现：")
                    for f in findings:
                        doc.add_paragraph(str(f), style="List Bullet")
                
                # 3. 气体浓度情况
                gases = extraction_result.get("gas_concentration", [])
                if gases:
                    doc.add_heading("3. 气体浓度监测", level=2)
                    latest_gas = gases[-1]
                    doc.add_paragraph(f"检测点数：{data_count} 次")
                    doc.add_paragraph(f"最新监测时间：{latest_gas.get('check_time')}")
                    doc.add_paragraph(f"平均氧气(O2)：{latest_gas.get('avg_o2', 0)}%")
                    doc.add_paragraph(f"平均磷化氢(PH3)：{latest_gas.get('avg_ph3', 0)} ppm")
                    doc.add_paragraph(f"平均二氧化碳(CO2)：{latest_gas.get('avg_co2', 0)}%")

                # 4. 趋势预测 (未来3天)
                try:
                    p = self.short_term_prediction(silo_id, prediction_days=3, start_time=report_start_time.strftime("%Y-%m-%d %H:%M:%S"), end_time=report_end_time.strftime("%Y-%m-%d %H:%M:%S"))
                    if p.get("status") != "no_data":
                        doc.add_heading("4. 趋势预测 (未来3天)", level=2)
                        doc.add_paragraph(f"预测趋势：{p.get('trend')}")
                        doc.add_paragraph(f"预计平均粮温将达到：{p.get('predicted_avg_temp')}°C")
                        doc.add_paragraph(f"风险预警：{p.get('risk_assessment')}")
                except Exception:
                    pass

                # 5. 综合分析与储藏建议 (LLM)
                try:
                    query_text = f"请对{silo_id}号仓在 {report_start_time.strftime('%Y-%m-%d')} 到 {report_end_time.strftime('%Y-%m-%d')} 期间的粮情进行综合分析。如果数据量（当前为{data_count}条）不足以支撑深入判断，请务必诚恳指出数据量过小的限制，不要进行无端推测。"
                    reasoning_result = self.llm_reasoning(
                        query=query_text,
                        context={
                            "analysis_result": a,
                            "warehouse_info": warehouse_info,
                            "latest_gas": gases[-1] if gases else None,
                            "prediction": p if 'p' in locals() else None,
                            "data_points_count": data_count,
                            "silo_id": silo_id,
                            "report_type": report_type,
                            "time_range": f"{report_start_time.strftime('%Y-%m-%d')} - {report_end_time.strftime('%Y-%m-%d')}"
                        }
                    )
                    
                    if reasoning_result.get("response"):
                        doc.add_heading("5. 专家诊断结论", level=2)
                        paragraphs = reasoning_result["response"].split("\n")
                        for para in paragraphs:
                            if para.strip(): doc.add_paragraph(para.strip())
                    
                    if reasoning_result.get("reasoning"):
                        doc.add_heading("6. 分析决策依据", level=2)
                        reasoning_paragraphs = reasoning_result["reasoning"].split("\n")
                        for para in reasoning_paragraphs:
                            if para.strip(): doc.add_paragraph(para.strip())
                except Exception as e:
                    logger.warning(f"LLM reasoning failed: {e}")
            except Exception as e:
                logger.warning(f"Detailed analysis section failed: {e}")

        doc.save(out_path)
        return {"report_id": report_id, "report_type": report_type, "silo_ids": silo_ids, "file_path": str(out_path), "status": "generated"}

