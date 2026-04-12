from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
doc.add_heading('Grain Agent V008 接口与业务规范说明书', 0)

# 1. 接入仓房列表接口
doc.add_heading('1. 接入仓房列表接口', level=1)
doc.add_paragraph('为了让智能体能够识别并映射用户的模糊指令（如 "P1仓"），系统新增了接入仓房列表接口。')
table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '字段名'
hdr_cells[1].text = '类型'
hdr_cells[2].text = '说明'
data = [
    ('house_code', 'String', '24位 WMS 系统唯一编码'),
    ('house_name', 'String', '仓房完整名称（如：西北库区 P1）'),
    ('short_name', 'String', '英文短名称或代号（如：P1）')
]
for field, type_, desc in data:
    row_cells = table.add_row().cells
    row_cells[0].text = field
    row_cells[1].text = type_
    row_cells[2].text = desc

# 2. API 访问字符串拼接规范
doc.add_heading('2. API 访问字符串拼接规范', level=1)
doc.add_paragraph('智能体在调用 WMS 接口时，必须遵循以下拼接规律：')
doc.add_paragraph('1) URL 结构：BASE_URL + /api/wms/{类别}/{动作}', style='List Bullet')
doc.add_paragraph('2) 强制参数：所有数据请求接口（粮温、气体等）必须携带 house_code 参数。', style='List Bullet')
doc.add_paragraph('3) 时间编码：start_time 和 end_time 参数必须包含空格（格式为 YYYY-MM-DD HH:MM:SS），并在拼接时进行标准 URL 编码。', style='List Bullet')
doc.add_paragraph('示例：', style='Normal')
doc.add_paragraph('http://.../grain/temperature?house_code=...&start_time=2026-01-01%2000:00:00', style='Quote')

# 3. 数据稀疏性熔断规则
doc.add_heading('3. 数据稀疏性熔断规则', level=1)
doc.add_paragraph('为保证分析结果的准确性，当历史数据点不足时，系统将触发熔断保护并给出友好提示：')
p = doc.add_paragraph()
p.add_run('1) 绘图熔断：').bold = True
p.add_run('当检测记录次数 <= 1 时，禁止生成三温图、两湿图或趋势图。')
p = doc.add_paragraph()
p.add_run('2) 预测熔断：').bold = True
p.add_run('当检测记录次数 <= 2 时，禁止生成短期趋势预测分析。')

doc.save('V008/Grain_Agent_V008_Spec.docx')
print("Successfully created V008/Grain_Agent_V008_Spec.docx")
