from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_api_section(doc, title, endpoint, description, params, example_json):
    doc.add_heading(title, level=2)
    p = doc.add_paragraph()
    p.add_run('端点 (Endpoint): ').bold = True
    p.add_run(endpoint)
    
    p = doc.add_paragraph()
    p.add_run('描述: ').bold = True
    p.add_run(description)
    
    if params:
        p = doc.add_paragraph()
        p.add_run('必要参数: ').bold = True
        p.add_run(', '.join(params))
        
    doc.add_paragraph('JSON 示例响应:', style='Normal').runs[0].bold = True
    code_block = doc.add_paragraph(example_json)
    code_block.style = 'Quote'

doc = Document()
doc.add_heading('Grain Agent V008 - WMS 接口对接技术规范', 0)

doc.add_heading('1. 通用说明', level=1)
doc.add_paragraph('本规范定义了粮情分析智能体与 WMS 系统对接的 4 个核心 API。WMS 侧工程师应确保接口返回的数据结构符合本规范要求。')
doc.add_paragraph('请求方式: GET', style='List Bullet')
doc.add_paragraph('时间格式: YYYY-MM-DD HH:MM:SS (URL 拼接需转义)', style='List Bullet')
doc.add_paragraph('数据格式: JSON', style='List Bullet')

doc.add_heading('2. 接口详解', level=1)

# API 1
add_api_section(
    doc, 
    '2.1 接入仓房列表接口 (Warehouse List)', 
    '/api/wms/warehouse/list',
    '获取当前所有授权智能体访问的仓房清单。智能体通过 short_name 识别用户指令并映射到 house_code。',
    None,
    '[\n  {\n    "house_code": "91620702MADKWU312X01001",\n    "house_name": "西北库区 P1",\n    "short_name": "P1"\n  }\n]'
)

# API 2
add_api_section(
    doc, 
    '2.2 仓房基本信息查询 (Warehouse Info)', 
    '/api/wms/warehouse/info',
    '查询指定仓房的库点名称、粮性、品种等静态属性。',
    ['house_code'],
    '{\n  "house_code": "91620702MADKWU312X01001",\n  "house_name": "西北库区 P1",\n  "depot_name": "中央储备粮西北库",\n  "grain_nature": "储备粮",\n  "variety": "小麦"\n}'
)

# API 3
add_api_section(
    doc, 
    '2.3 粮温数据查询 (Grain Temperature)', 
    '/api/wms/grain/temperature',
    '查询指定时间段内的历史温湿度记录。',
    ['house_code', 'start_time', 'end_time'],
    '[\n  {\n    "house_code": "91620702MADKWU312X01001",\n    "check_time": "2026-01-02 10:00:00",\n    "indoor_temp": 22.5,\n    "indoor_humidity": 45.0,\n    "outdoor_temp": 15.0,\n    "avg_temp": 20.1,\n    "temp_values": "20.1,1,1,1|20.5,1,1,2|..."\n  }\n]'
)

# API 4
add_api_section(
    doc, 
    '2.4 气体浓度查询 (Gas Concentration)', 
    '/api/wms/gas/concentration',
    '查询指定时间段内的气体（O2, PH3, CO2）浓度记录。',
    ['house_code', 'start_time', 'end_time'],
    '[\n  {\n    "house_code": "91620702MADKWU312X01012",\n    "check_time": "2026-01-02 10:30:00",\n    "avg_o2": 20.1,\n    "avg_ph3": 150,\n    "avg_co2": 0.5\n  }\n]'
)

doc.add_heading('3. 业务逻辑提醒', level=1)
doc.add_paragraph('数据稀疏性熔断: 绘图至少需 2 个点，预测至少需 3 个点。若数据量不足，智能体将返回友好提示而非错误。', style='List Bullet')
doc.add_paragraph('URL 拼接: 智能体会将时间空格编码为 %20。例: ?start_time=2026-01-01%2000:00:00', style='List Bullet')

doc.save('V008/WMS_API_Integration_Spec_V008.docx')
print("WMS API Spec Word document created successfully.")
