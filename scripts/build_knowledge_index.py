"""
知识库索引构建脚本

读取 data/knowledge/ 目录下的 .md / .txt / .docx / .pdf 文件，
分块 → Embedding → 写入 ChromaDB。

用法:
    cd V008
    python -m scripts.build_knowledge_index
"""

import os, sys, re, logging
from pathlib import Path

# 确保项目根目录在 sys.path 中
BASE_DIR = Path(__file__).resolve().parent.parent  # V008/
sys.path.insert(0, str(BASE_DIR))

from dotenv import load_dotenv
load_dotenv(BASE_DIR / ".env", override=True)

logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s")
logger = logging.getLogger(__name__)

# --------------- 配置 ---------------
KNOWLEDGE_DIR = BASE_DIR / "data" / "knowledge"
CHROMA_DIR = BASE_DIR / "data" / "chroma_db"
COLLECTION_NAME = "grain_knowledge"
CHUNK_MAX_CHARS = 500
CHUNK_OVERLAP = 50
EMBEDDING_BATCH_SIZE = 10


# --------------- 文件读取 ---------------
def _table_to_markdown(table) -> str:
    """将表格对象转为 Markdown 表格文本（通用辅助函数）"""
    if not table:
        return ""
    rows = []
    for row in table:
        cells = [str(c).strip().replace("\n", " ") if c else "" for c in row]
        rows.append("| " + " | ".join(cells) + " |")
    if len(rows) >= 1:
        # 在表头后插入分隔行
        sep = "| " + " | ".join(["---"] * len(table[0])) + " |"
        rows.insert(1, sep)
    return "\n".join(rows)


def read_docx(fpath: Path) -> str:
    """读取 .docx 文件，按文档顺序提取段落和表格，表格转为 Markdown 格式"""
    from docx import Document
    from docx.oxml.ns import qn

    doc = Document(str(fpath))
    parts: list[str] = []

    # 按 XML 顺序遍历 body 子元素，保持段落与表格的原始顺序
    for element in doc.element.body:
        tag = element.tag

        if tag.endswith("}p"):  # 段落
            # 从段落 element 反查 Paragraph 对象
            for p in doc.paragraphs:
                if p._element is element:
                    if p.text.strip():
                        parts.append(p.text)
                    break

        elif tag.endswith("}tbl"):  # 表格
            for tbl in doc.tables:
                if tbl._element is element:
                    rows = []
                    for row in tbl.rows:
                        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                        rows.append(cells)
                    md = _table_to_markdown(rows)
                    if md:
                        parts.append(md)
                    break

    return "\n\n".join(parts)


def read_pdf(fpath: Path) -> str:
    """读取 .pdf 文件，分别提取文本和表格，表格转为 Markdown 格式"""
    import pdfplumber

    pages: list[str] = []
    with pdfplumber.open(str(fpath)) as pdf:
        for page in pdf.pages:
            page_parts: list[str] = []

            # 提取表格区域，用于从正文中排除
            tables = page.extract_tables()
            table_bboxes = []
            if page.find_tables():
                table_bboxes = [t.bbox for t in page.find_tables()]

            # 提取非表格区域的文本
            if table_bboxes:
                # 裁剪掉表格区域后提取文本
                filtered_page = page
                for bbox in table_bboxes:
                    # 用 crop 的补集方式：提取表格区域外的文本
                    pass
                # 简化处理：直接提取全页文本（表格文本会重复，但下面会追加结构化表格）
                text = page.extract_text()
                if text:
                    page_parts.append(text.strip())
            else:
                text = page.extract_text()
                if text:
                    page_parts.append(text.strip())

            # 追加结构化表格
            for table_data in tables:
                md = _table_to_markdown(table_data)
                if md:
                    page_parts.append(md)

            if page_parts:
                pages.append("\n\n".join(page_parts))

    return "\n\n".join(pages)


def read_file(fpath: Path) -> str:
    """根据扩展名选择读取方式"""
    ext = fpath.suffix.lower()
    if ext == ".docx":
        return read_docx(fpath)
    elif ext == ".pdf":
        return read_pdf(fpath)
    else:  # .md, .txt
        return fpath.read_text(encoding="utf-8")


# --------------- 分块 ---------------
def split_by_markdown_heading(text: str) -> list[dict]:
    """按 Markdown 标题拆分，返回 [{title, content}]"""
    sections: list[dict] = []
    current_title = ""
    current_lines: list[str] = []

    for line in text.splitlines():
        if re.match(r"^#{1,4}\s+", line):
            # 保存上一段
            if current_lines:
                sections.append({"title": current_title, "content": "\n".join(current_lines).strip()})
            current_title = line.lstrip("#").strip()
            current_lines = []
        else:
            current_lines.append(line)

    # 最后一段
    if current_lines:
        sections.append({"title": current_title, "content": "\n".join(current_lines).strip()})

    return sections


def split_long_section(text: str, max_chars: int = CHUNK_MAX_CHARS, overlap: int = CHUNK_OVERLAP) -> list[str]:
    """超长段落按段落边界二次切分，带重叠"""
    if len(text) <= max_chars:
        return [text]

    paragraphs = re.split(r"\n{2,}", text)
    chunks: list[str] = []
    current_chunk = ""

    for para in paragraphs:
        if len(current_chunk) + len(para) + 2 > max_chars and current_chunk:
            chunks.append(current_chunk.strip())
            # 重叠：取末尾 overlap 个字符作为下一段开头
            current_chunk = current_chunk[-overlap:] + "\n\n" + para
        else:
            current_chunk = (current_chunk + "\n\n" + para) if current_chunk else para

    if current_chunk.strip():
        chunks.append(current_chunk.strip())

    return chunks


def load_and_chunk(knowledge_dir: Path) -> list[dict]:
    """读取所有知识文件并分块，返回 [{id, text, metadata}]"""
    all_chunks: list[dict] = []
    supported_exts = ("*.md", "*.txt", "*.docx", "*.pdf")
    files = []
    for ext in supported_exts:
        files.extend(knowledge_dir.glob(ext))

    if not files:
        logger.warning(f"知识目录 {knowledge_dir} 中未找到支持的文件（.md/.txt/.docx/.pdf）")
        return all_chunks

    for fpath in sorted(files):
        logger.info(f"处理文件: {fpath.name}")
        try:
            text = read_file(fpath)
        except Exception as e:
            logger.error(f"读取文件失败 {fpath.name}: {e}")
            continue
        source = fpath.name

        sections = split_by_markdown_heading(text)
        if not sections:
            sections = [{"title": "", "content": text}]

        for sec in sections:
            if not sec["content"]:
                continue
            sub_chunks = split_long_section(sec["content"])
            for i, chunk_text in enumerate(sub_chunks):
                chunk_id = f"{source}::{sec['title']}::{i}"
                all_chunks.append({
                    "id": chunk_id,
                    "text": chunk_text,
                    "metadata": {
                        "source": source,
                        "title": sec["title"],
                    },
                })

    logger.info(f"共生成 {len(all_chunks)} 个文本块")
    return all_chunks


# --------------- Embedding ---------------
def embed_texts(texts: list[str], client, model: str, dimensions: int) -> list[list[float]]:
    """批量 embedding，每批 EMBEDDING_BATCH_SIZE 条"""
    all_embeddings: list[list[float]] = []
    for i in range(0, len(texts), EMBEDDING_BATCH_SIZE):
        batch = texts[i : i + EMBEDDING_BATCH_SIZE]
        resp = client.embeddings.create(model=model, input=batch, dimensions=dimensions)
        all_embeddings.extend([d.embedding for d in resp.data])
        logger.info(f"  Embedding 批次 {i // EMBEDDING_BATCH_SIZE + 1} 完成 ({len(batch)} 条)")
    return all_embeddings


# --------------- 主流程 ---------------
def main():
    from app.core.config import settings

    # 1. 检查知识目录
    if not KNOWLEDGE_DIR.exists():
        logger.error(f"知识目录不存在: {KNOWLEDGE_DIR}")
        logger.info("请创建 data/knowledge/ 目录并放入 .md / .txt / .docx / .pdf 知识文件")
        sys.exit(1)

    # 2. 分块
    chunks = load_and_chunk(KNOWLEDGE_DIR)
    if not chunks:
        logger.error("未生成任何文本块，请检查知识文件内容")
        sys.exit(1)

    # 3. Embedding
    from openai import OpenAI
    embed_client = OpenAI(
        api_key=settings.EMBEDDING_API_KEY or settings.DASHSCOPE_API_KEY,
        base_url=settings.EMBEDDING_BASE_URL or settings.LLM_BASE_URL,
    )

    texts = [c["text"] for c in chunks]
    logger.info(f"开始生成 Embedding（模型: {settings.EMBEDDING_MODEL}，维度: {settings.EMBEDDING_DIMENSIONS}）...")
    embeddings = embed_texts(texts, embed_client, settings.EMBEDDING_MODEL, settings.EMBEDDING_DIMENSIONS)

    # 4. 写入 ChromaDB
    CHROMA_DIR.mkdir(parents=True, exist_ok=True)

    import chromadb
    client = chromadb.PersistentClient(path=str(CHROMA_DIR))

    # 删除旧 collection（如果存在）再重建
    try:
        client.delete_collection(COLLECTION_NAME)
        logger.info(f"已删除旧 collection '{COLLECTION_NAME}'")
    except Exception:
        pass

    collection = client.create_collection(
        name=COLLECTION_NAME,
        metadata={"hnsw:space": "cosine"},
    )

    # 批量写入
    ids = [c["id"] for c in chunks]
    documents = texts
    metadatas = [c["metadata"] for c in chunks]

    collection.add(
        ids=ids,
        documents=documents,
        embeddings=embeddings,
        metadatas=metadatas,
    )

    logger.info(f"索引构建完成！共写入 {collection.count()} 个文本块到 {CHROMA_DIR}")
    logger.info(f"Collection: {COLLECTION_NAME}")


if __name__ == "__main__":
    main()
